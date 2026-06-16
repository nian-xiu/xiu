from __future__ import annotations

from datetime import datetime
from pathlib import Path
import zipfile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
NON_PROJECT = ROOT / "Non-project-content"

REPORT_PATH = DOCS / "心怡商城毕业设计风格项目报告.docx"
RESUME_TXT_PATH = NON_PROJECT / "心怡商城项目经历.txt"
ENV_TXT_PATH = NON_PROJECT / "项目运行环境.txt"

CH_FONT = "Microsoft YaHei"
FONT_CANDIDATES = [
    Path("C:/Windows/Fonts/msyh.ttc"),
    Path("C:/Windows/Fonts/simhei.ttf"),
    Path("C:/Windows/Fonts/simsun.ttc"),
]
FONT_PATH = next((path for path in FONT_CANDIDATES if path.exists()), None)


def image_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if FONT_PATH is not None:
        return ImageFont.truetype(str(FONT_PATH), size=size, index=0)
    return ImageFont.load_default()


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = CH_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CH_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.5):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def apply_table_style(table, widths: list[float]):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = tc_pr.find(qn("w:tcMar"))
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for name in ["top", "bottom", "start", "end"]:
                node = mar.find(qn(f"w:{name}"))
                if node is None:
                    node = OxmlElement(f"w:{name}")
                    mar.append(node)
                node.set(qn("w:w"), "120")
                node.set(qn("w:type"), "dxa")


def setup_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = CH_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CH_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    heading_specs = [
        ("Title", 22, "173F4F"),
        ("Heading 1", 16, "173F4F"),
        ("Heading 2", 13, "285B63"),
        ("Heading 3", 11.5, "6F4C1E"),
    ]
    for style_name, size, color in heading_specs:
        style = styles[style_name]
        style.font.name = CH_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CH_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("心怡商城网上商城系统项目报告")
    set_run_font(header_run, size=9, color="666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("毕业设计论文大纲参考版")
    set_run_font(footer_run, size=9, color="777777")
    return document


def add_paragraph(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_heading(document: Document, text: str, level: int = 1):
    return document.add_paragraph(text, style=f"Heading {level}")


def add_caption(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color="666666")


def add_image(document: Document, path: Path, caption: str, width: float = 6.35):
    if path.exists():
        document.add_picture(str(path), width=Inches(width))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(document, caption)


def add_bullets(document: Document, items: list[str]):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = document.add_table(rows=1, cols=len(headers))
    apply_table_style(table, widths)
    for index, header in enumerate(headers):
        shade_cell(table.rows[0].cells[index], "E8F0F2")
        set_cell_text(table.rows[0].cells[index], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            set_cell_text(cells[index], text)
    return table


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_sequence(title: str, participants: list[str], messages: list[tuple[int, int, str]], filename: str):
    width = 1800
    height = max(980, 250 + len(messages) * 82)
    image = Image.new("RGB", (width, height), "#fbfaf6")
    draw = ImageDraw.Draw(image)
    title_font = image_font(38, True)
    label_font = image_font(24, True)
    msg_font = image_font(22)
    small_font = image_font(18)

    draw.rounded_rectangle([32, 28, width - 32, height - 28], radius=28, fill="#fffdfa", outline="#d7c59d", width=4)
    draw.text((72, 60), title, fill="#173F4F", font=title_font)
    draw.text((72, 108), "浏览器请求经过 Controller、Service、Mapper 到 MySQL 的主要调用链", fill="#6a6256", font=small_font)

    top = 185
    bottom = height - 90
    left_margin = 120
    right_margin = 120
    gap = (width - left_margin - right_margin) / (len(participants) - 1)
    xs = [left_margin + int(index * gap) for index in range(len(participants))]

    for x, label in zip(xs, participants):
        draw.rounded_rectangle([x - 96, top - 48, x + 96, top + 16], radius=15, fill="#285B63", outline="#173F4F", width=2)
        label_lines = wrap_text(draw, label, label_font, 168)
        y = top - 38 if len(label_lines) == 1 else top - 43
        for line in label_lines[:2]:
            text_width = draw.textlength(line, font=label_font)
            draw.text((x - text_width / 2, y), line, fill="#ffffff", font=label_font)
            y += 25
        draw.line([x, top + 18, x, bottom], fill="#c5b796", width=3)

    y = top + 88
    for index, (source, target, label) in enumerate(messages, start=1):
        x1 = xs[source]
        x2 = xs[target]
        color = "#285B63" if index % 2 else "#B38A3C"
        draw.line([x1, y, x2, y], fill=color, width=4)
        arrow = 14 if x2 >= x1 else -14
        draw.polygon([(x2, y), (x2 - arrow, y - 9), (x2 - arrow, y + 9)], fill=color)
        label_text = f"{index}. {label}"
        max_label_width = max(180, abs(x2 - x1) - 30)
        label_lines = wrap_text(draw, label_text, msg_font, max_label_width)
        text_width = max(draw.textlength(line, font=msg_font) for line in label_lines)
        box_height = 28 * len(label_lines) + 12
        box_x = (x1 + x2) / 2 - text_width / 2 - 12
        box_y = y - box_height - 10
        draw.rounded_rectangle([box_x, box_y, box_x + text_width + 24, box_y + box_height], radius=9, fill="#fffdfa", outline="#eadfc5")
        text_y = box_y + 6
        for line in label_lines:
            draw.text((box_x + 12, text_y), line, fill="#25221d", font=msg_font)
            text_y += 28
        y += 82

    image.save(ASSETS / filename)


def build_sequence_diagrams():
    ASSETS.mkdir(parents=True, exist_ok=True)
    roles = ["用户/管理员", "Thymeleaf页面", "Controller", "Service", "Mapper", "MySQL"]
    diagrams = [
        (
            "用户注册与登录功能时序图",
            [
                (0, 1, "填写注册或登录表单"),
                (1, 2, "提交 /register 或 /login"),
                (2, 3, "表单校验、密码摘要比对"),
                (3, 4, "查询或写入用户记录"),
                (4, 5, "执行 users 表 SQL"),
                (3, 2, "返回用户对象或错误信息"),
                (2, 1, "写入 Session 并跳转首页/后台"),
            ],
            "seq-auth-new.png",
        ),
        (
            "首页、商品搜索与筛选功能时序图",
            [
                (0, 1, "进入首页或输入关键词"),
                (1, 2, "GET /、/products、/search"),
                (2, 3, "构造 ProductFilter 并读取分类"),
                (3, 4, "调用分类和商品查询接口"),
                (4, 5, "按分类、价格、排序和黑名单过滤"),
                (5, 4, "返回商品集合"),
                (2, 1, "渲染首页、列表页和筛选结果"),
            ],
            "seq-catalog-new.png",
        ),
        (
            "商品详情、收藏与黑名单功能时序图",
            [
                (0, 1, "查看详情或点击收藏/拉黑"),
                (1, 2, "GET /products/{id} 或 POST /favorites/toggle"),
                (2, 3, "读取商品、收藏状态和黑名单状态"),
                (3, 4, "调用 Product/Favorite/Blacklist Mapper"),
                (4, 5, "查询商品或增删用户行为记录"),
                (3, 2, "返回页面模型"),
                (2, 1, "刷新详情页或收藏页"),
            ],
            "seq-favorite-new.png",
        ),
        (
            "购物车管理功能时序图",
            [
                (0, 1, "加入购物车、改数量或删除"),
                (1, 2, "POST /cart/add、/cart/{id}/update、/remove"),
                (2, 3, "校验登录用户、商品状态和库存"),
                (3, 4, "查询购物车已有项"),
                (4, 5, "写入或更新 cart_items"),
                (3, 2, "汇总数量与小计"),
                (2, 1, "跳转购物车或商品详情"),
            ],
            "seq-cart-new.png",
        ),
        (
            "立即购买与订单结算功能时序图",
            [
                (0, 1, "选择地址、支付方式、优惠券或金币"),
                (1, 2, "POST /checkout 或 /checkout/buy-now"),
                (2, 3, "校验微信确认、读取地址和商品明细"),
                (3, 4, "开启事务创建订单"),
                (4, 5, "写 orders、order_items 并扣库存"),
                (3, 4, "扣金币、核销优惠券、清空购物车"),
                (2, 1, "跳转订单详情"),
            ],
            "seq-order-new.png",
        ),
        (
            "订单查询、退款、收货与客服留言时序图",
            [
                (0, 1, "查看订单、申请退款、确认收货或留言"),
                (1, 2, "GET /orders/{id} 或 POST /orders/{id}/..."),
                (2, 3, "按用户读取订单详情并校验状态"),
                (3, 4, "查询订单明细或写客服消息"),
                (4, 5, "访问 orders、order_items、service_messages"),
                (3, 2, "返回状态变更结果"),
                (2, 1, "刷新订单详情页"),
            ],
            "seq-user-order-new.png",
        ),
        (
            "收货地址管理功能时序图",
            [
                (0, 1, "新增、编辑、设为默认或删除地址"),
                (1, 2, "POST /addresses 或 /addresses/{id}"),
                (2, 3, "校验 AddressForm 与当前用户"),
                (3, 4, "保存地址并处理默认地址互斥"),
                (4, 5, "更新 addresses 表"),
                (3, 2, "返回保存结果"),
                (2, 1, "刷新地址列表或结算页"),
            ],
            "seq-address-new.png",
        ),
        (
            "福利活动、背包、站内信和兑换码功能时序图",
            [
                (0, 1, "领取活动、兑换 CDKEY 或领取邮件"),
                (1, 2, "POST /activity/{id}/claim、/redeem、/mail/{id}/claim"),
                (2, 3, "校验活动状态、名额、过期和重复领取"),
                (3, 4, "记录领取并发放金币或优惠券"),
                (4, 5, "写 activity_claims、user_coupons、reward_mails"),
                (3, 2, "返回领取结果"),
                (2, 1, "刷新活动页、背包页或站内信"),
            ],
            "seq-reward-new.png",
        ),
        (
            "个人资料与密码设置功能时序图",
            [
                (0, 1, "修改昵称、手机号、邮箱或密码"),
                (1, 2, "POST /settings/profile 或 /settings/password"),
                (2, 3, "执行 Bean Validation 和旧密码校验"),
                (3, 4, "更新用户资料或密码摘要"),
                (4, 5, "写 users 表"),
                (3, 2, "返回更新后的用户"),
                (2, 1, "更新 Session 并刷新设置页"),
            ],
            "seq-settings-new.png",
        ),
        (
            "用户客服与后台回复功能时序图",
            [
                (0, 1, "用户留言或管理员回复"),
                (1, 2, "POST /service/messages 或 /admin/service/reply"),
                (2, 3, "校验消息内容、用户与订单归属"),
                (3, 4, "创建用户消息或管理员回复"),
                (4, 5, "写 service_messages 并维护未读标记"),
                (3, 2, "返回回复结果"),
                (2, 1, "刷新客服中心或后台客服页"),
            ],
            "seq-service-new.png",
        ),
        (
            "后台商品管理功能时序图",
            [
                (0, 1, "新增、编辑、上传图片或删除商品"),
                (1, 2, "POST /admin/products 或 /admin/products/{id}"),
                (2, 3, "校验管理员权限和 ProductForm"),
                (3, 3, "保存上传图片到 uploads/products"),
                (3, 4, "调用商品保存接口"),
                (4, 5, "insert/update/delete products"),
                (2, 1, "跳转后台商品列表"),
            ],
            "seq-admin-product-new.png",
        ),
        (
            "后台订单处理功能时序图",
            [
                (0, 1, "筛选订单、查看详情或修改状态"),
                (1, 2, "GET /admin/orders 或 POST /admin/orders/{id}/status"),
                (2, 3, "校验 ADMIN 权限并读取订单"),
                (3, 3, "校验状态流转合法性"),
                (3, 4, "查询或更新订单状态"),
                (4, 5, "访问 orders、order_items、service_messages"),
                (2, 1, "渲染订单列表或详情页"),
            ],
            "seq-admin-order-new.png",
        ),
        (
            "后台用户与权限管理功能时序图",
            [
                (0, 1, "查看用户或启用/停用账号"),
                (1, 2, "GET /admin/users 或 POST /admin/users/{id}/status"),
                (2, 3, "AuthInterceptor 校验 ADMIN"),
                (3, 4, "读取用户列表或更新状态"),
                (4, 5, "访问 users 表"),
                (3, 2, "返回结果"),
                (2, 1, "刷新后台用户列表"),
            ],
            "seq-admin-user-new.png",
        ),
        (
            "后台公告管理功能时序图",
            [
                (0, 1, "新建、编辑、发布、下线或置顶公告"),
                (1, 2, "POST /admin/announcements/save、/status、/pinned"),
                (2, 3, "校验表单与管理员身份"),
                (3, 4, "保存公告、切换状态或置顶"),
                (4, 5, "写 announcements 表"),
                (3, 2, "返回操作结果"),
                (2, 1, "刷新公告管理页"),
            ],
            "seq-admin-announcement-new.png",
        ),
        (
            "后台福利运营功能时序图",
            [
                (0, 1, "创建活动、群发邮件、生成兑换码"),
                (1, 2, "POST /admin/activities/campaigns、mails、codes"),
                (2, 3, "校验奖励类型、额度、有效期和名额"),
                (3, 4, "写入活动、邮件或兑换码"),
                (4, 5, "访问 activity_campaigns、reward_mails、coupon_codes"),
                (3, 2, "返回生成结果"),
                (2, 1, "刷新福利运营页"),
            ],
            "seq-admin-reward-new.png",
        ),
    ]
    for title, messages, filename in diagrams:
        draw_sequence(title, roles, messages, filename)
    return [filename for _, _, filename in diagrams]


def build_report(diagram_files: list[str]):
    document = setup_document()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("心怡商城网上商城系统项目报告")
    set_run_font(run, size=22, bold=True, color="173F4F")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("参考毕业设计论文大纲框架：需求分析、系统设计、详细设计、系统实现与总结")
    set_run_font(run, size=11.5, color="666666")

    date_line = document.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_line.add_run(f"生成日期：{datetime.now():%Y年%m月%d日}")
    set_run_font(run, size=10, color="777777")

    add_heading(document, "摘 要", 1)
    add_paragraph(
        document,
        "心怡商城是一套基于 SSM 分层思想实现的网上商城系统。项目使用 Spring Boot 3 进行工程组织，"
        "以 Spring MVC 处理 Web 请求，以 Spring Service 承担业务编排和事务控制，以 MyBatis XML Mapper 完成数据访问，"
        "并结合 Thymeleaf、MySQL、Maven 和 Undertow 构成完整运行环境。系统围绕电商业务闭环展开，"
        "实现了商品浏览、搜索筛选、商品详情、用户注册登录、收藏与黑名单、购物车、立即购买、购物车结算、"
        "地址管理、订单查询、退款与确认收货、客服消息、签到、活动福利、背包优惠券、站内信、兑换码，"
        "以及后台商品、订单、用户、公告和福利运营管理等功能。"
    )
    add_paragraph(
        document,
        "本文档按照毕业设计论文的写作框架，对项目需求、总体设计、数据库设计、详细功能设计、"
        "SSM 框架配置过程、系统实现效果和开发心得进行说明，其中详细设计部分重点给出各功能的调用逻辑与时序图。"
    )
    add_paragraph(document, "关键词：SSM；Spring MVC；MyBatis；网上商城；Thymeleaf；MySQL")

    add_heading(document, "目 录", 1)
    for entry in [
        "一、需求分析",
        "二、系统总体设计",
        "三、数据库设计",
        "四、详细设计与功能时序图",
        "五、SSM 框架下系统的配置过程",
        "六、系统实现",
        "七、测试与运行说明",
        "八、总结与心得",
    ]:
        add_paragraph(document, entry)

    document.add_page_break()
    add_heading(document, "一、需求分析", 1)
    add_heading(document, "1.1 项目背景", 2)
    add_paragraph(
        document,
        "随着线上购物场景逐渐成为日常消费的重要方式，网上商城系统需要同时支持商品展示、用户行为记录、"
        "订单处理、营销活动和后台运营。心怡商城面向课程设计和项目展示场景，重点体现 Java Web 项目的分层架构、"
        "数据库驱动能力、前后台协同能力和可运行交付能力。项目不依赖外部 Tomcat，而是使用内嵌 Undertow，"
        "适合在本地快速启动、演示和调试。"
    )
    add_heading(document, "1.2 用户角色分析", 2)
    add_table(
        document,
        ["角色", "主要目标", "核心功能", "访问限制"],
        [
            ["游客", "快速了解商品信息", "首页、分类、搜索、商品详情、公告查看", "不可加入购物车或下单"],
            ["普通用户", "完成购买与售后流程", "注册登录、收藏、购物车、地址、下单、订单、客服、福利中心", "需登录"],
            ["管理员", "维护商城数据与运营活动", "后台仪表盘、商品、订单、用户、公告、福利运营、客服回复", "需 ADMIN 角色"],
        ],
        [1.2, 2.0, 2.5, 1.1],
    )
    add_heading(document, "1.3 功能需求", 2)
    add_bullets(
        document,
        [
            "商品展示需求：支持首页推荐商品、最新商品、分类导航、关键词搜索、价格筛选、排序和商品详情展示。",
            "用户账户需求：支持注册、登录、退出、个人资料修改、密码修改，并通过 Session 维护登录态。",
            "购物流程需求：支持收藏商品、加入购物车、调整数量、删除购物车项、立即购买、购物车结算和微信扫码模拟确认。",
            "订单业务需求：支持订单生成、库存扣减、优惠券或金币抵扣、订单详情、退款申请、确认收货和客服留言。",
            "会员运营需求：支持签到、活动领取、站内奖励邮件、背包优惠券、兑换码和新人首单券。",
            "后台管理需求：支持商品 CRUD、商品图片上传、订单状态处理、用户启停、公告发布、福利活动与兑换码管理。",
        ],
    )
    add_heading(document, "1.4 非功能需求", 2)
    add_bullets(
        document,
        [
            "可维护性：采用 Controller、Service、Mapper、Domain、DTO、Form、Template 分层目录，降低模块耦合。",
            "一致性：下单、扣库存、优惠核销、金币扣减和购物车清空在事务内完成，避免局部成功。",
            "安全性：通过登录拦截器保护用户中心和后台路径，通过 CSRF 拦截器为表单请求提供基础防护。",
            "易部署性：通过 Maven 构建可运行 JAR，配置集中在 application.yml，数据库脚本位于 resources/db。",
            "展示性：使用 Thymeleaf 页面、统一布局片段和静态资源，使课程答辩或简历展示具备直观界面。",
        ],
    )

    add_heading(document, "二、系统总体设计", 1)
    add_heading(document, "2.1 技术架构", 2)
    add_paragraph(
        document,
        "系统采用浏览器/服务器架构。浏览器端通过 Thymeleaf 渲染页面并结合少量原生 JavaScript 完成交互；"
        "服务端由 Spring Boot 启动，Spring MVC 负责路由映射和模型返回，Service 层负责核心业务规则，"
        "MyBatis Mapper 层负责 SQL 映射，MySQL 负责持久化数据。整体仍然保持传统 SSM 的分层思想，"
        "同时借助 Spring Boot 自动配置降低 XML 配置和外部容器部署成本。"
    )
    add_table(
        document,
        ["层次", "对应目录/类", "职责说明"],
        [
            ["表示层", "templates、static、fragments/layout.html", "展示商城前台、用户中心和后台管理页面"],
            ["控制层", "controller 包", "接收请求、绑定参数、处理校验结果、选择视图或重定向"],
            ["业务层", "service 包", "实现商品、购物车、订单、用户、活动、客服等业务流程"],
            ["持久层", "mapper 接口与 mapper/*.xml", "封装 MyBatis SQL 查询、插入、更新和删除"],
            ["数据层", "MySQL、schema.sql、data.sql", "存储用户、商品、订单、活动、公告、客服消息等数据"],
        ],
        [1.35, 2.35, 2.8],
    )
    add_heading(document, "2.2 模块结构设计", 2)
    add_paragraph(
        document,
        "项目模块按照业务边界划分：商城浏览模块由 ShopController 与 CatalogService 负责；购物车模块由 CartController "
        "与 CartService 负责；订单模块由 OrderController 与 OrderService 负责；用户账户模块由 AuthController、"
        "SettingsController 与 UserService 负责；福利运营模块由 RewardCenterController、AdminRewardController "
        "与 RewardCenterService 负责；后台管理模块由 AdminController、AdminAnnouncementController 等控制器组成。"
    )
    add_table(
        document,
        ["模块", "主要入口", "主要服务", "说明"],
        [
            ["商城浏览", "/、/products、/products/{id}", "CatalogService", "分类、推荐、搜索、详情和黑名单过滤"],
            ["账户与个人中心", "/login、/register、/settings", "UserService", "注册登录、资料维护和密码修改"],
            ["购物车与结算", "/cart、/checkout", "CartService、OrderService", "购物车汇总、立即购买、订单生成"],
            ["订单与客服", "/orders、/service", "OrderService、ServiceMessageService", "订单状态、退款、收货、客服留言"],
            ["福利中心", "/activity、/backpack、/mail、/redeem", "RewardCenterService", "活动、优惠券、邮件、兑换码"],
            ["后台管理", "/admin/**", "AdminService、CatalogService、RewardCenterService", "运营仪表盘和各类管理功能"],
        ],
        [1.3, 1.7, 2.0, 1.8],
    )

    add_heading(document, "三、数据库设计", 1)
    add_heading(document, "3.1 数据库总体说明", 2)
    add_paragraph(
        document,
        "数据库名为 ssm_shop，初始化脚本位于 src/main/resources/db/schema.sql 和 data.sql。"
        "系统围绕用户、商品、购物车、订单、地址和运营活动建立表结构，其中 orders 与 order_items 构成订单主从关系，"
        "activity_campaigns、activity_claims、reward_mails、coupon_codes、user_coupons 等表支持福利与优惠能力。"
    )
    add_table(
        document,
        ["表名", "用途"],
        [
            ["users", "用户账号、昵称、角色、状态、金币等信息"],
            ["categories / products", "商品分类、商品价格、库存、销量、封面和上下架状态"],
            ["cart_items / favorites / product_blacklist", "购物车、收藏与商品黑名单行为数据"],
            ["addresses", "用户收货地址和默认地址标记"],
            ["orders / order_items", "订单主表、订单明细、金额、优惠、状态和物流时间"],
            ["announcements", "商城公告、置顶状态、发布时间和过期时间"],
            ["service_messages", "用户和管理员客服消息、未读状态"],
            ["activity_campaigns / activity_claims", "公开福利活动和用户领取记录"],
            ["reward_mails / coupon_codes / user_coupons", "站内奖励邮件、兑换码和用户优惠券背包"],
        ],
        [2.2, 4.1],
    )
    add_heading(document, "3.2 关键数据关系", 2)
    add_paragraph(
        document,
        "users 与 cart_items、favorites、addresses、orders、user_coupons、reward_mails 等表形成一对多关系；"
        "products 与 order_items、cart_items、favorites 关联；orders 与 order_items 形成一对多主从关系。"
        "下单时系统会从购物车或立即购买摘要生成订单，写入订单主表和明细表，同时扣减 products.stock、增加 products.sales，"
        "并根据用户选择核销 user_coupons 或扣减 users.coins。"
    )

    add_heading(document, "四、详细设计与功能时序图", 1)
    add_paragraph(
        document,
        "详细设计部分按照项目的主要功能逐一说明。各功能基本遵循“用户或管理员触发页面操作、Controller 接收请求、"
        "Service 进行业务校验和事务处理、Mapper 访问数据库、页面显示结果”的调用链。"
    )

    details = [
        ("4.1 用户注册与登录", "注册流程通过 RegisterForm 接收用户名、密码、昵称、手机号和邮箱等信息，UserService 负责检查用户名唯一性并生成密码摘要，注册成功后将 currentUser 写入 Session。登录流程通过 LoginForm 接收账号密码，UserService 查询 users 表并比对密码摘要，账号停用或密码错误时返回错误信息，管理员登录后跳转后台首页。"),
        ("4.2 首页、商品搜索与筛选", "首页读取分类、推荐商品、最新商品和签到状态；商品列表通过 ProductFilter 支持分类、关键词、价格区间和排序组合查询。对于已登录用户，商品查询会结合黑名单过滤，让被屏蔽的商品不再出现在普通列表中。"),
        ("4.3 商品详情、收藏与黑名单", "商品详情页展示商品基础信息、价格、库存、销量、评分和描述。登录用户可收藏或取消收藏，也可将商品加入黑名单。收藏和黑名单都以用户 ID 与商品 ID 为联合行为记录，便于后续查询和过滤。"),
        ("4.4 购物车管理", "购物车模块在添加商品时校验商品是否存在、是否上架以及库存是否充足；若购物车中已有同一商品，则累加数量并限制不超过库存。更新数量与删除操作都带有 userId 条件，避免越权修改其他用户的购物车项。"),
        ("4.5 立即购买与订单结算", "结算模块支持购物车结算和立即购买两种模式。OrderService 会读取地址、商品明细、支付方式、金币或优惠券选择，在事务中生成订单号、写订单主表和明细表、扣库存、增加销量、扣金币、核销优惠券，并在购物车结算场景下清空购物车。"),
        ("4.6 订单查询、退款、收货与客服留言", "用户订单页按当前用户读取订单列表，订单详情页展示订单主信息、商品明细和关联客服消息。退款申请和确认收货都需要校验当前订单状态，客服留言会写入 service_messages 表并同步给后台。"),
        ("4.7 收货地址管理", "地址模块支持新增、编辑、删除和默认地址设置。保存默认地址时，AddressService 会先清除该用户原默认地址，再保存新的默认地址，从而保证同一用户最多只有一个默认地址。"),
        ("4.8 福利活动、背包、站内信和兑换码", "福利中心支持公开活动领取、站内奖励邮件领取、兑换码兑换和优惠券背包管理。RewardCenterService 会校验活动状态、领取窗口、名额、重复领取和过期条件，再发放金币或优惠券。"),
        ("4.9 个人资料与密码设置", "设置模块支持修改昵称、手机号、邮箱和登录密码。修改个人资料后会更新 Session 中的 currentUser；修改密码时需要校验旧密码，并保存新的密码摘要。"),
        ("4.10 用户客服与后台回复", "用户可以在客服中心或订单详情中提交咨询内容，后台管理员可在客服页或订单详情中回复。消息表记录发送角色、关联订单和用户/管理员未读状态，便于前后台同步提醒。"),
        ("4.11 后台商品管理", "后台商品管理支持查询、新增、编辑、删除和封面图片上传。上传图片由 ProductImageStorageService 存入 uploads/products，商品信息由 CatalogService 写入 products 表。"),
        ("4.12 后台订单处理", "后台订单列表支持按状态筛选，详情页展示订单和客服消息。管理员修改订单状态时，OrderService 会根据当前状态校验流转是否合法，避免已完成、已退款或已取消订单被错误回退。"),
        ("4.13 后台用户与权限管理", "后台用户管理展示所有用户，并支持启用或停用账号。AuthInterceptor 对 /admin 路径进行 ADMIN 角色校验，普通用户无法进入后台。"),
        ("4.14 后台公告管理", "公告管理支持公告新增、编辑、发布、下线和置顶切换。公告可设置分类、发布时间、过期时间和关联活动，用于首页或公告页面展示运营信息。"),
        ("4.15 后台福利运营", "福利运营支持创建公开活动、活动启停、活动克隆、定向或群发奖励邮件、生成兑换码、兑换码启停和删除。该模块使管理员能够模拟真实电商运营中的优惠发放和会员激励。"),
    ]
    for (heading, text), filename in zip(details, diagram_files, strict=True):
        add_heading(document, heading, 2)
        add_paragraph(document, text)
        add_image(document, ASSETS / filename, f"图 {heading.split('.')[0]}-{heading.split('.')[1].split()[0]} {heading[4:]}功能时序图", width=6.35)

    add_heading(document, "五、SSM 框架下系统的配置过程", 1)
    add_heading(document, "5.1 Maven 依赖配置", 2)
    add_paragraph(
        document,
        "项目 pom.xml 以 spring-boot-starter-parent 3.5.7 为父工程，Java 版本为 21。核心依赖包括 "
        "spring-boot-starter-web、spring-boot-starter-thymeleaf、spring-boot-starter-validation、"
        "spring-security-crypto、mybatis-spring-boot-starter、mysql-connector-j 和 spring-boot-starter-test。"
        "为了使用 Undertow，项目在 web starter 中排除 spring-boot-starter-tomcat，并引入 spring-boot-starter-undertow。"
    )
    add_heading(document, "5.2 Spring MVC 与拦截器配置", 2)
    add_paragraph(
        document,
        "WebConfig 实现 WebMvcConfigurer，注册 AuthInterceptor 和 CsrfInterceptor。AuthInterceptor 保护购物车、结算、订单、"
        "地址、收藏、客服、福利中心、设置和后台等路径；当访问 /admin 开头路径时，还会检查当前用户角色是否为 ADMIN。"
        "CsrfInterceptor 为表单请求提供 CSRF 校验，并排除静态资源、上传资源、错误页和微信确认页。"
    )
    add_heading(document, "5.3 MyBatis 配置", 2)
    add_paragraph(
        document,
        "启动类 SsmShopApplication 使用 @MapperScan(\"com.example.ssmshop.mapper\") 扫描 Mapper 接口。"
        "application.yml 中配置 mybatis.mapper-locations 为 classpath:mapper/*.xml，type-aliases-package 为 "
        "com.example.ssmshop.domain，并开启 map-underscore-to-camel-case，使数据库下划线字段可以映射到 Java 驼峰属性。"
    )
    add_heading(document, "5.4 数据源与数据库初始化配置", 2)
    add_paragraph(
        document,
        "application.yml 中配置 MySQL 驱动、连接地址、用户名、密码和 HikariCP 连接池参数。数据库地址默认为 "
        "jdbc:mysql://localhost:3306/ssm_shop，账号默认为 root，密码默认为 123456。项目的 SQL 初始化模式设置为 never，"
        "由 DataSeeder 在首次启动时执行 schema.sql 和 data.sql，并通过 app_meta.schema_initialized 标记避免重复初始化。"
    )
    add_heading(document, "5.5 Thymeleaf、静态资源与上传目录配置", 2)
    add_paragraph(
        document,
        "Thymeleaf 模板位于 src/main/resources/templates，静态资源位于 src/main/resources/static。"
        "WebConfig 将 /uploads/products/** 映射到 app.upload.product-image-dir 指定的本地目录，默认目录为 uploads/products，"
        "用于后台商品封面上传和前台商品图片展示。"
    )
    add_heading(document, "5.6 Undertow 与运行参数配置", 2)
    add_paragraph(
        document,
        "server.port 默认为 8080，Undertow 的 IO 线程数配置为 4，Worker 线程数配置为 16，并开启 HTML、CSS、JavaScript、"
        "JSON 和 SVG 等资源压缩。Windows 环境下若默认临时目录权限导致 Undertow docbase 初始化失败，可在启动时指定 "
        "-Djava.io.tmpdir=项目目录/tmp，项目中的 tools/run-app.ps1 已考虑该运行方式。"
    )

    add_heading(document, "六、系统实现", 1)
    add_heading(document, "6.1 前台商城实现", 2)
    add_paragraph(
        document,
        "前台页面包含首页、商品列表、商品详情、购物车、结算页、微信确认页和错误页等。首页展示品牌入口、分类导航、"
        "推荐商品和最新商品；商品列表页支持筛选和排序；商品详情页提供收藏、拉黑、立即购买和加入购物车入口；"
        "结算页整合地址、商品清单、金币、优惠券和支付方式。"
    )
    for filename, caption in [
        ("screenshot-home.png", "图 6-1 商城首页展示"),
        ("screenshot-products.png", "图 6-2 商品列表与筛选展示"),
        ("screenshot-detail.png", "图 6-3 商品详情页展示"),
        ("screenshot-login.png", "图 6-4 用户登录页展示"),
    ]:
        add_image(document, ASSETS / filename, caption, width=6.35)
    add_heading(document, "6.2 用户中心实现", 2)
    add_paragraph(
        document,
        "用户中心实现订单、地址、收藏、黑名单、客服、活动、背包、站内信和个人设置等页面。"
        "这些页面均依赖当前 Session 中的用户 ID 查询数据，保证用户只能访问自己的订单、地址、优惠券和客服消息。"
    )
    add_heading(document, "6.3 后台管理实现", 2)
    add_paragraph(
        document,
        "后台管理以 /admin 为统一入口。仪表盘展示商品数量、订单数量、用户数量和销售额等统计信息；"
        "商品管理支持图片上传和商品信息维护；订单管理支持状态筛选、详情查看和状态变更；"
        "用户管理支持账号启停；公告管理和福利运营用于模拟真实商城的内容与营销管理。"
    )
    add_image(document, ASSETS / "screenshot-admin.png", "图 6-5 后台仪表盘展示", width=6.35)

    add_heading(document, "七、测试与运行说明", 1)
    add_paragraph(
        document,
        "项目可通过 Maven 进行编译打包，生成 target/ssm-shop-1.0.0.jar 后以 java -jar 方式运行。"
        "运行前需要准备 JDK 21、Maven、MySQL，并创建 ssm_shop 数据库。启动成功后访问 http://localhost:8080/ "
        "进入商城首页，访问 http://localhost:8080/admin 进入后台。演示账号包括普通用户 customer / 123456 和管理员 admin / admin123。"
    )
    add_bullets(
        document,
        [
            "功能验证：逐项访问首页、商品列表、登录注册、购物车、结算、订单、地址、活动、客服和后台页面。",
            "数据验证：检查下单后 orders、order_items、products.stock、products.sales、cart_items 等表变化是否符合预期。",
            "权限验证：未登录访问购物车、订单、地址和后台时应跳转登录页，普通用户访问后台应返回首页。",
            "异常验证：库存不足、优惠券不可用、订单状态非法流转、重复领取活动等情况应显示错误提示。",
        ],
    )

    add_heading(document, "八、总结与心得", 1)
    add_paragraph(
        document,
        "通过本项目的设计与实现，我对 SSM 分层架构、MVC 请求流程、MyBatis XML 映射、事务控制、权限拦截、"
        "数据库初始化和前后台页面协同有了更加系统的理解。与简单的增删改查项目相比，商城系统的难点在于多个业务对象之间存在强关联，"
        "例如购物车、库存、订单、优惠券、金币和用户地址必须在同一结算流程中协同工作。"
    )
    add_paragraph(
        document,
        "本项目的最大收获是认识到“清晰分层”和“事务边界”对后端系统十分重要。Controller 不应承载复杂业务规则，"
        "Service 层需要集中处理校验、状态流转和数据库写入顺序，Mapper 层则应保持 SQL 明确可追踪。"
        "同时，良好的页面组织和演示数据也会显著提升项目展示效果。后续若继续完善，可以加入真实支付沙箱、商品评价、分页、"
        "日志审计、更严格的密码加密策略、接口化前后端分离以及自动化测试，使系统更接近生产级电商项目。"
    )

    DOCS.mkdir(parents=True, exist_ok=True)
    document.save(REPORT_PATH)


def build_text_files():
    NON_PROJECT.mkdir(parents=True, exist_ok=True)
    resume_text = (
        "项目经历：心怡商城网上商城系统\n"
        "基于 Java、Spring Boot、Spring MVC、MyBatis、Thymeleaf、MySQL 和 Undertow 开发 B2C 网上商城系统，"
        "实现商品浏览、分类筛选、详情展示、注册登录、收藏、购物车、立即购买、地址管理、优惠券/金币抵扣、订单查询、退款申请、确认收货、客服留言、活动福利、站内信、兑换码和后台运营管理。"
        "本人负责需求拆分、模块设计、数据库建表、Mapper XML 编写、后端业务实现、页面联调和本地部署排障。"
        "后端采用控制层、业务层、持久层分层结构，在业务层通过事务完成下单、扣库存、写订单明细、核销优惠券、扣减金币和清空购物车，保证交易链路一致性；"
        "通过登录拦截器实现用户权限校验和后台管理员访问控制。"
        "后台完成商品增删改查、订单状态流转、用户启停、公告发布、活动邮件和兑换码管理。"
        "通过该项目提升了电商业务建模、SSM 分层开发、数据一致性处理和 Windows 本地调试能力。"
    )

    env_text = (
        "心怡商城项目运行环境\n\n"
        "1. 操作系统：Windows 10/11 或其他可运行 Java 21、Maven、MySQL 的系统。\n"
        "2. JDK：Java 21，pom.xml 中 java.version 配置为 21。\n"
        "3. 构建工具：Maven 3.8+，用于下载依赖、编译和打包项目。\n"
        "4. 数据库：MySQL 8.0/9.x，需提前创建数据库 ssm_shop。\n"
        "5. 数据库连接：src/main/resources/application.yml 默认配置为 jdbc:mysql://localhost:3306/ssm_shop，用户名 root，密码 123456，可按本机环境修改。\n"
        "6. Web 容器：项目使用内嵌 Undertow，不需要单独安装 Tomcat。\n"
        "7. 主要框架/依赖：Spring Boot 3.5.7、Spring MVC、Thymeleaf、Spring Validation、Spring Security Crypto、MyBatis Spring Boot Starter 3.0.5、MySQL Connector/J。\n"
        "8. 端口：默认访问端口为 8080，商城首页为 http://localhost:8080/，后台入口为 http://localhost:8080/admin。\n"
        "9. 数据初始化：建库后首次启动由 DataSeeder 执行 src/main/resources/db/schema.sql 和 data.sql，写入表结构与演示数据。\n"
        "10. 推荐启动：在项目根目录执行 mvn package -DskipTests 后运行 java -Djava.io.tmpdir=tmp -jar target/ssm-shop-1.0.0.jar；Windows 下也可执行 powershell -ExecutionPolicy Bypass -File tools/run-app.ps1。\n"
        "11. 演示账号：普通用户 customer / 123456；管理员 admin / admin123。\n"
        "12. 上传目录：商品图片默认保存到 uploads/products，对应配置项 app.upload.product-image-dir。\n"
    )
    RESUME_TXT_PATH.write_text(resume_text, encoding="utf-8", newline="\n")
    ENV_TXT_PATH.write_text(env_text, encoding="utf-8", newline="\n")


def validate_docx(path: Path):
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        missing = required - names
        if missing:
            raise RuntimeError(f"DOCX 缺少必要文件: {', '.join(sorted(missing))}")
        document_xml = archive.read("word/document.xml")
        if "心怡商城".encode("utf-8") not in document_xml:
            raise RuntimeError("DOCX 正文校验失败，未找到项目名称")


def main():
    diagram_files = build_sequence_diagrams()
    build_report(diagram_files)
    build_text_files()
    validate_docx(REPORT_PATH)
    print(REPORT_PATH)
    print(RESUME_TXT_PATH)
    print(ENV_TXT_PATH)


if __name__ == "__main__":
    main()
