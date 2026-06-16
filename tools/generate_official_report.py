from __future__ import annotations

import re
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "report-v2-assets"
OUT = DOCS / "心怡商城课程设计报告_正式版.docx"
OFFICIAL_ASSETS = DOCS / "report-official-assets"
DIAGRAMS = OFFICIAL_ASSETS / "diagrams"

BODY_FONT = "宋体"
HEADING_FONT = "黑体"
ASCII_FONT = "Times New Roman"
DIAGRAM_FONT_PATH = next(
    (
        p
        for p in [
            Path("C:/Windows/Fonts/simhei.ttf"),
            Path("C:/Windows/Fonts/msyh.ttc"),
            Path("C:/Windows/Fonts/simsun.ttc"),
        ]
        if p.exists()
    ),
    None,
)

INK = "#1f2933"
MUTED = "#637381"
GREEN = "#17695c"
GREEN_DARK = "#0f4c43"
BLUE = "#2d6f8f"
ORANGE = "#b45f2a"
PURPLE = "#6b4aa5"
BORDER = "#ccd9d5"
SOFT = "#f5f8f7"
TABLE_FILL = "EEF6F3"


def font(size: int):
    if DIAGRAM_FONT_PATH:
        return ImageFont.truetype(str(DIAGRAM_FONT_PATH), size=size)
    return ImageFont.load_default()


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None, name: str | None = None):
    east_asia = name or BODY_FONT
    run.font.name = ASCII_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, *, size: float = 10.5, bold: bool = False, color: str | None = None, name: str = BODY_FONT):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color, name=name)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color="777777")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run = paragraph.add_run()
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_sep)
    page_run._r.append(fld_text)
    page_run._r.append(fld_end)
    set_run_font(page_run, size=9, color="777777")
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color="777777")


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = ASCII_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.first_line_indent = Inches(0.29)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(3)

    for style_name, size, color, before, after in [
        ("Heading 1", 15, "17483F", 8, 5),
        ("Heading 2", 12.5, "285B63", 5, 3),
        ("Heading 3", 11, "6D4C15", 4, 2),
    ]:
        style = styles[style_name]
        style.font.name = ASCII_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.12

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("心怡商城网上商城系统课程设计报告")
    set_run_font(run, size=9, color="777777")
    add_page_number(section.footer.paragraphs[0])
    return doc


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_paragraph_font(p, size=15 if level == 1 else 12.5 if level == 2 else 11, bold=True, name=HEADING_FONT)
    return p


def add_para(doc: Document, text: str, *, first_line: bool = True, align=None, after: float = 3):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.29) if first_line else None
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_small(doc: Document, text: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, after: float = 2):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.space_after = Pt(after)
    p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=9.5, color="555555")
    return p


def page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, *, bold: bool = False, size: float = 9.4, color: str | None = None, align=None):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, TABLE_FILL)
        set_cell(cell, text, bold=True, size=9.6, color="17483F", align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell(cells[i], text, size=9.2)
            cells[i].width = Inches(widths[i])
    return table


def caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, size=9, color="666666")


def add_image(doc: Document, image: Path, cap: str, width: float):
    if not image.exists():
        raise FileNotFoundError(image)
    doc.add_picture(str(image), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap)


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt, width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        if draw.textlength(current + char, font=fnt) <= width:
            current += char
        else:
            if current:
                lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def rounded(draw, xy, *, fill: str, outline: str = BORDER, radius: int = 18, width: int = 2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_arrow(draw, x1: int, y1: int, x2: int, y2: int, *, color: str = GREEN, width: int = 4, dashed: bool = False):
    if dashed:
        gap = 16
        dist = abs(x2 - x1)
        direction = 1 if x2 >= x1 else -1
        current = x1
        while abs(current - x1) < dist:
            nx = current + direction * min(gap, dist - abs(current - x1))
            draw.line([current, y1, nx, y2], fill=color, width=width)
            current = nx + direction * gap
    else:
        draw.line([x1, y1, x2, y2], fill=color, width=width)
    direction = 1 if x2 >= x1 else -1
    draw.polygon([(x2, y2), (x2 - direction * 13, y2 - 8), (x2 - direction * 13, y2 + 8)], fill=color)


def make_architecture_diagram() -> Path:
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    path = DIAGRAMS / "architecture.png"
    img = Image.new("RGB", (1600, 760), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((65, 42), "心怡商城系统总体架构", fill=GREEN_DARK, font=font(36))
    d.text((65, 88), "采用 MVC 分层组织页面、控制、业务和持久化代码，降低页面逻辑与数据库操作之间的耦合。", fill=MUTED, font=font(22))

    boxes = [
        (80, 170, 350, 300, "用户界面", "Thymeleaf 页面、CSS、JS、表单", BLUE),
        (430, 170, 700, 300, "Controller", "路由映射、参数校验、视图模型", GREEN),
        (780, 170, 1050, 300, "Service", "业务规则、事务边界、状态流转", ORANGE),
        (430, 430, 700, 560, "Mapper", "MyBatis XML、动态查询、表关联", GREEN),
        (780, 430, 1050, 560, "MySQL", "商品、订单、用户、活动与客服数据", BLUE),
        (1170, 300, 1480, 430, "运行基础", "Spring Boot、Undertow、Maven、上传目录", PURPLE),
    ]
    for x1, y1, x2, y2, title, sub, accent in boxes:
        rounded(d, [x1, y1, x2, y2], fill=SOFT, outline=BORDER, radius=20, width=2)
        d.rounded_rectangle([x1, y1, x1 + 12, y2], radius=5, fill=accent)
        d.text((x1 + 30, y1 + 22), title, fill=INK, font=font(27))
        for idx, line in enumerate(wrap(d, sub, font(19), x2 - x1 - 45)[:2]):
            d.text((x1 + 30, y1 + 66 + idx * 25), line, fill=MUTED, font=font(19))
    draw_arrow(d, 350, 235, 430, 235, color=BLUE)
    draw_arrow(d, 700, 235, 780, 235, color=GREEN)
    draw_arrow(d, 915, 300, 565, 430, color=ORANGE)
    draw_arrow(d, 700, 495, 780, 495, color=GREEN)
    draw_arrow(d, 1050, 495, 1170, 365, color=BLUE)
    d.text((80, 650), "设计重点：Controller 不直接处理复杂业务；Service 统一处理交易、活动、客服等规则；Mapper 负责 SQL 与实体映射。", fill=INK, font=font(23))
    img.save(path)
    return path


def make_sequence_diagram(title: str, messages: list[tuple[int, int, str, bool]], filename: str) -> Path:
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    roles = ["页面", "Controller", "Service", "Mapper", "MySQL"]
    xs = [145, 455, 765, 1075, 1385]
    img = Image.new("RGB", (1600, 760), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((60, 35), title, fill=GREEN_DARK, font=font(34))
    d.text((60, 78), "按页面请求、控制层调度、业务处理、持久化访问和结果返回的顺序描述。", fill=MUTED, font=font(20))
    top = 158
    bottom = 705
    for x, role in zip(xs, roles):
        rounded(d, [x - 92, top - 44, x + 92, top + 14], fill="#eff7f4", outline="#bed8d0", radius=16)
        tw = d.textlength(role, font=font(23))
        d.text((x - tw / 2, top - 29), role, fill=GREEN_DARK, font=font(23))
        d.line([x, top + 25, x, bottom], fill="#dce7e4", width=3)

    y = 220
    step_h = max(58, min(72, 430 // max(1, len(messages))))
    for index, (src, dst, label, ret) in enumerate(messages, start=1):
        x1, x2 = xs[src], xs[dst]
        color = MUTED if ret else GREEN
        draw_arrow(d, x1 + (18 if dst > src else -18), y, x2 - (18 if dst > src else -18), y, color=color, dashed=ret, width=3 if ret else 4)
        label_text = f"{index}. {label}"
        label_font = font(18)
        lines = wrap(d, label_text, label_font, abs(x2 - x1) - 35)[:2]
        tx = min(x1, x2) + 28
        ty = y - 26 if len(lines) == 1 else y - 34
        for line_no, line in enumerate(lines):
            d.text((tx, ty + line_no * 22), line, fill=INK, font=label_font)
        y += step_h

    path = DIAGRAMS / filename
    img.save(path)
    return path


def make_config_diagram() -> Path:
    messages = [
        (0, 1, "启动应用并读取 application.yml", False),
        (1, 2, "装配 Spring MVC、拦截器和资源映射", False),
        (2, 3, "初始化 MyBatis Mapper 与连接池", False),
        (3, 4, "连接 ssm_shop 数据库", False),
        (2, 4, "DataSeeder 检查表结构和迁移标记", False),
        (4, 2, "返回初始化结果，进入可访问状态", True),
    ]
    return make_sequence_diagram("SSM 框架配置与启动流程", messages, "seq-config.png")


def paste_cover(draw: ImageDraw.ImageDraw, canvas: Image.Image, image_path: Path, xy: tuple[int, int], size: tuple[int, int], label: str):
    src = Image.open(image_path).convert("RGB")
    src_ratio = src.width / src.height
    target_w, target_h = size
    target_ratio = target_w / target_h
    if src_ratio > target_ratio:
        crop_w = int(src.height * target_ratio)
        x0 = (src.width - crop_w) // 2
        src = src.crop((x0, 0, x0 + crop_w, src.height))
    else:
        crop_h = int(src.width / target_ratio)
        y0 = max(0, (src.height - crop_h) // 3)
        src = src.crop((0, y0, src.width, y0 + crop_h))
    src = src.resize(size, Image.Resampling.LANCZOS)
    x, y = xy
    shadow = Image.new("RGBA", (target_w + 24, target_h + 24), (0, 0, 0, 0))
    sd = ImageDraw.Draw(shadow)
    sd.rounded_rectangle([12, 12, target_w + 12, target_h + 12], radius=18, fill=(20, 40, 35, 28))
    canvas.paste(shadow.convert("RGB"), (x - 12, y - 12))
    canvas.paste(src, (x, y))
    draw.rounded_rectangle([x, y, x + target_w, y + target_h], radius=16, outline="#c7d8d2", width=3)
    draw.rounded_rectangle([x + 18, y + 18, x + 142, y + 54], radius=14, fill="#ffffff", outline="#d3e1dc", width=1)
    draw.text((x + 36, y + 26), label, fill=GREEN_DARK, font=font(20))


def make_screenshot_board(title: str, left: Path, right: Path, left_label: str, right_label: str, filename: str) -> Path:
    OFFICIAL_ASSETS.mkdir(parents=True, exist_ok=True)
    path = OFFICIAL_ASSETS / filename
    img = Image.new("RGB", (1600, 980), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((70, 50), title, fill=GREEN_DARK, font=font(36))
    d.text((70, 95), "页面来源于当前项目运行界面，按前台、用户中心和后台运营场景分组展示。", fill=MUTED, font=font(21))
    paste_cover(d, img, left, (70, 175), (700, 560), left_label)
    paste_cover(d, img, right, (830, 175), (700, 560), right_label)
    d.rounded_rectangle([70, 810, 1530, 900], radius=18, fill="#f6f8f7", outline="#d9e5e1", width=2)
    d.text((105, 838), "实现说明：页面复用统一导航与表单片段，前后台通过相同业务服务读写数据，保证展示结果与后台维护状态一致。", fill=INK, font=font(23))
    img.save(path)
    return path


def create_diagrams() -> dict[str, Path]:
    diagrams = {"architecture": make_architecture_diagram(), "config": make_config_diagram()}
    specs = {
        "auth": (
            "用户注册与登录时序图",
            [
                (0, 1, "提交注册或登录表单", False),
                (1, 2, "校验参数、验证码与跳转地址", False),
                (2, 3, "查询用户或写入新用户", False),
                (3, 4, "执行用户表读写", False),
                (2, 1, "返回登录结果并写入 Session", True),
                (1, 0, "跳转到首页或原访问页面", True),
            ],
        ),
        "catalog": (
            "商品浏览与检索时序图",
            [
                (0, 1, "访问首页、分类页或搜索页", False),
                (1, 2, "组装分类、关键词、价格和排序条件", False),
                (2, 3, "调用商品与分类查询", False),
                (3, 4, "读取商品、库存、评分和分类数据", False),
                (2, 1, "返回商品列表与筛选状态", True),
                (1, 0, "渲染商品卡片和分页结果", True),
            ],
        ),
        "favorite": (
            "详情、收藏与黑名单时序图",
            [
                (0, 1, "打开商品详情或点击收藏", False),
                (1, 2, "检查登录状态和商品状态", False),
                (2, 3, "查询详情、收藏记录或黑名单记录", False),
                (3, 4, "读写 favorites、product_blacklists", False),
                (2, 1, "返回详情数据或操作结果", True),
                (1, 0, "刷新详情页、收藏页或黑名单页", True),
            ],
        ),
        "cart": (
            "购物车维护时序图",
            [
                (0, 1, "添加商品、修改数量或移除条目", False),
                (1, 2, "校验用户、商品和数量", False),
                (2, 3, "读取购物车与商品库存", False),
                (3, 4, "更新 cart_items 记录", False),
                (2, 1, "重新计算小计和合计", True),
                (1, 0, "返回购物车页面", True),
            ],
        ),
        "checkout": (
            "结算下单时序图",
            [
                (0, 1, "进入结算页并选择地址、优惠和支付方式", False),
                (1, 2, "校验库存、优惠券、金币和微信确认状态", False),
                (2, 3, "写入订单、明细并扣减库存", False),
                (3, 4, "更新 orders、order_items、products", False),
                (2, 1, "返回订单编号和状态", True),
                (1, 0, "跳转订单详情页", True),
            ],
        ),
        "order_service": (
            "订单售后与客服时序图",
            [
                (0, 1, "查看订单、确认收货、退款或留言", False),
                (1, 2, "校验订单归属和可操作状态", False),
                (2, 3, "读取订单明细或写入客服消息", False),
                (3, 4, "更新订单状态和 service_messages", False),
                (2, 1, "返回状态变化与消息列表", True),
                (1, 0, "展示订单进度和客服回复", True),
            ],
        ),
        "profile": (
            "地址管理与个人设置时序图",
            [
                (0, 1, "提交地址、资料或密码表单", False),
                (1, 2, "校验必填项、手机号和原密码", False),
                (2, 3, "调用用户与地址持久化接口", False),
                (3, 4, "更新 users、addresses", False),
                (2, 1, "返回保存结果", True),
                (1, 0, "回到地址页或设置页", True),
            ],
        ),
        "activity": (
            "活动领取与兑换码时序图",
            [
                (0, 1, "进入活动页或提交兑换码", False),
                (1, 2, "判断活动时间、库存、领取次数和码状态", False),
                (2, 3, "写入领取记录并发放金币或优惠券", False),
                (3, 4, "更新 activity_claims、user_coupons", False),
                (2, 1, "返回奖励名称和领取状态", True),
                (1, 0, "刷新活动页或背包页", True),
            ],
        ),
        "mail": (
            "站内信与背包时序图",
            [
                (0, 1, "打开邮件或背包页面", False),
                (1, 2, "标记已读并筛选可领取奖励", False),
                (2, 3, "读取邮件、优惠券和有效期", False),
                (3, 4, "更新 reward_mails、user_coupons", False),
                (2, 1, "返回未读数、可用券和清理结果", True),
                (1, 0, "展示邮件列表和优惠券卡片", True),
            ],
        ),
        "admin_product": (
            "后台商品管理时序图",
            [
                (0, 1, "进入商品管理并提交新增或编辑表单", False),
                (1, 2, "校验管理员权限、表单和图片文件", False),
                (2, 3, "保存图片并写入商品信息", False),
                (3, 4, "更新 products 与分类关联数据", False),
                (2, 1, "返回商品列表和提示信息", True),
                (1, 0, "展示后台商品管理页面", True),
            ],
        ),
        "admin_order": (
            "后台订单与用户管理时序图",
            [
                (0, 1, "查看订单、修改状态或启停用户", False),
                (1, 2, "校验管理员身份和目标记录", False),
                (2, 3, "调用订单或用户管理接口", False),
                (3, 4, "更新 orders、users、service_messages", False),
                (2, 1, "返回操作结果与统计数据", True),
                (1, 0, "刷新仪表盘、订单页或用户页", True),
            ],
        ),
        "admin_activity": (
            "后台公告与运营活动时序图",
            [
                (0, 1, "发布公告、活动、邮件或兑换码", False),
                (1, 2, "校验标题、奖励、数量和时间范围", False),
                (2, 3, "写入运营配置和发放目标", False),
                (3, 4, "更新 announcements、campaigns、coupon_codes", False),
                (2, 1, "返回启用、暂停、克隆或删除结果", True),
                (1, 0, "展示运营中心最新状态", True),
            ],
        ),
    }
    for key, (title, messages) in specs.items():
        diagrams[key] = make_sequence_diagram(title, messages, f"seq-{key}.png")
    diagrams["front_board"] = make_screenshot_board(
        "前台商城页面展示",
        ASSETS / "current-home.png",
        ASSETS / "current-products.png",
        "首页",
        "商品列表",
        "board-front.png",
    )
    diagrams["user_board"] = make_screenshot_board(
        "用户中心页面展示",
        ASSETS / "current-activity.png",
        ASSETS / "current-backpack.png",
        "活动中心",
        "用户背包",
        "board-user.png",
    )
    diagrams["admin_board"] = make_screenshot_board(
        "后台运营页面展示",
        ASSETS / "current-admin-dashboard.png",
        ASSETS / "current-admin-rewards.png",
        "仪表盘",
        "活动运营",
        "board-admin.png",
    )
    return diagrams


def cover(doc: Document):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run("心怡商城网上商城系统课程设计报告")
    set_run_font(r, size=25, bold=True, color="17483F", name=HEADING_FONT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run("基于 Spring MVC + Spring + MyBatis 的 SSM 商城系统")
    set_run_font(r, size=14, color="444444", name=BODY_FONT)
    for _ in range(7):
        doc.add_paragraph()
    rows = [
        ("课程名称", "JavaEE 课程设计"),
        ("项目名称", "心怡商城"),
        ("技术路线", "Spring Boot、Spring MVC、MyBatis、Thymeleaf、MySQL、Undertow"),
        ("完成时间", "2026 年 6 月"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.6)
        cells[1].width = Inches(4.5)
        set_cell(cells[0], label, bold=True, size=10.5, color="17483F", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(cells[1], value, size=10.5)
    page_break(doc)


def front_matter(doc: Document):
    add_heading(doc, "摘  要", 1)
    add_para(
        doc,
        "心怡商城是一个面向普通用户和后台管理员的网上商城系统。系统以 Spring MVC、Spring、MyBatis 为主要开发框架，结合 Thymeleaf、MySQL 和 Undertow，实现商品浏览、登录注册、收藏、购物车、地址、下单、订单售后、活动福利、站内信、公告以及后台运营管理等功能。项目重点放在商城业务流程的完整性和页面操作的连贯性上，既能完成前台用户从选品到下单的主要流程，也能让管理员维护商品、订单、用户和运营活动。本文从需求分析、系统设计、详细设计、SSM 配置过程和系统实现几个方面进行说明，并通过时序图描述主要功能的交互过程。",
    )
    add_para(
        doc,
        "关键词：网上商城；SSM；Spring MVC；MyBatis；Thymeleaf；订单管理",
    )
    add_heading(doc, "目  录", 1)
    toc = [
        "一、需求分析",
        "二、系统设计",
        "三、详细设计与时序图",
        "四、SSM 框架配置过程",
        "五、系统实现与页面展示",
        "六、总结与心得",
    ]
    for item in toc:
        add_small(doc, item, after=3)
    page_break(doc)


def requirements_and_design(doc: Document, diagrams: dict[str, Path]):
    add_heading(doc, "一、需求分析", 1)
    add_heading(doc, "1.1 项目背景", 2)
    add_para(
        doc,
        "网上商城系统的核心是把商品展示、用户购买和后台维护串联成稳定流程。心怡商城不是只完成单个页面，而是形成可连续操作的课程设计项目：未登录用户可以浏览商品和公告，普通用户登录后可以加入购物车、管理地址、参与活动、使用优惠券并查看订单，管理员负责商品、订单、用户、公告和福利活动维护。",
    )
    add_heading(doc, "1.2 用户角色与功能需求", 2)
    add_table(
        doc,
        ["角色", "主要需求", "说明"],
        [
            ["游客", "浏览首页、商品列表、商品详情、公告", "不需要登录即可查看公开内容。"],
            ["普通用户", "注册登录、收藏、购物车、下单、地址、订单、客服、活动、背包、邮件", "覆盖从选品到售后的完整购物流程。"],
            ["管理员", "仪表盘、商品维护、订单处理、用户启停、公告发布、活动和兑换码管理", "用于维护平台数据和处理运营事务。"],
        ],
        [1.1, 3.1, 2.6],
    )
    add_heading(doc, "1.3 非功能需求", 2)
    add_para(
        doc,
        "系统需要具备基本的安全性、易用性和可维护性。安全性体现在登录拦截、CSRF 校验、订单归属校验和文件上传校验；易用性体现在导航清晰、筛选方便和订单状态可见；可维护性体现在控制层、业务层和持久层职责分离。",
    )

    add_heading(doc, "二、系统设计", 1)
    add_heading(doc, "2.1 总体架构设计", 2)
    add_para(
        doc,
        "系统采用典型 MVC 分层。浏览器端由 Thymeleaf 页面和静态资源组成，Controller 接收请求并准备视图数据，Service 负责业务判断和事务处理，Mapper 通过 XML SQL 访问 MySQL。该结构使页面、业务和数据库操作相对独立，适合持续扩展。",
    )
    add_image(doc, diagrams["architecture"], "图 1 系统总体架构图", 6.25)
    add_heading(doc, "2.2 模块划分", 2)
    add_table(
        doc,
        ["模块", "包含功能", "关键类或资源"],
        [
            ["前台商城", "首页推荐、分类、搜索、详情、收藏、黑名单", "ShopController、CatalogService、FavoriteService"],
            ["交易流程", "购物车、结算、地址、优惠、订单、微信确认", "CartService、OrderService、AddressService"],
            ["用户中心", "订单、客服、活动、背包、邮件、个人设置", "RewardCenterService、ServiceMessageService、SettingsController"],
            ["后台管理", "商品、订单、用户、公告、运营活动和兑换码", "AdminController、AdminRewardController、AdminAnnouncementController"],
            ["系统配置", "拦截器、资源映射、数据初始化、MyBatis 配置", "WebConfig、DataSeeder、application.yml"],
        ],
        [1.2, 2.75, 2.85],
    )
    add_para(
        doc,
        "数据库设计围绕用户、商品、订单和运营数据展开。订单主表与明细表记录交易，活动、邮件和优惠券表记录奖励来源，公告表负责通知内容。通过字段关联和业务校验控制数据关系，既能满足页面查询，也方便后台追踪状态。",
    )
    page_break(doc)


def detailed_design(doc: Document, diagrams: dict[str, Path]):
    add_heading(doc, "三、详细设计与时序图", 1)
    add_para(
        doc,
        "详细设计主要说明各功能在运行时如何协作。以下时序图以页面、Controller、Service、Mapper 和 MySQL 为固定参与者，展示请求进入系统后的处理顺序。为控制篇幅，同一模块中相近的操作放在一张图内描述。",
    )
    groups = [
        ("3.1 账号与商品浏览设计", "账号模块负责注册、登录、退出和安全跳转；商品模块负责首页推荐、分类筛选、关键词搜索与详情展示。登录成功后用户信息写入 Session，后续需要身份的功能由拦截器统一判断。", [("auth", "图 2 用户注册与登录时序图"), ("catalog", "图 3 商品浏览与检索时序图")]),
        ("3.2 收藏、黑名单与购物车设计", "收藏功能用于保留感兴趣商品，黑名单用于屏蔽不希望继续看到的商品。购物车以用户为维度记录商品和数量，更新数量时同时校验库存，避免结算阶段出现明显的数据不一致。", [("favorite", "图 4 详情、收藏与黑名单时序图"), ("cart", "图 5 购物车维护时序图")]),
        ("3.3 结算、订单与售后设计", "结算时系统会读取购物车或立即购买商品、收货地址、金币和优惠券，并在创建订单时扣减库存。订单详情页不仅显示商品和物流状态，也提供确认收货、退款申请和客服留言入口。", [("checkout", "图 6 结算下单时序图"), ("order_service", "图 7 订单售后与客服时序图")]),
        ("3.4 用户资料、福利与消息设计", "用户资料和地址管理偏向基础数据维护；活动、兑换码、背包和站内信偏向运营奖励。活动领取需要判断时间、库存和重复领取，站内信支持单封领取和集中领取，背包用于查看可用与失效优惠券。", [("profile", "图 8 地址管理与个人设置时序图"), ("activity", "图 9 活动领取与兑换码时序图"), ("mail", "图 10 站内信与背包时序图")]),
        ("3.5 后台管理设计", "后台功能按维护对象划分。商品管理需要处理表单与图片上传，订单和用户管理需要控制状态变化，公告与活动管理则用于发布通知、配置奖励、克隆活动和启停兑换码。", [("admin_product", "图 11 后台商品管理时序图"), ("admin_order", "图 12 后台订单与用户管理时序图"), ("admin_activity", "图 13 后台公告与运营活动时序图")]),
    ]
    for idx, (heading, text, figs) in enumerate(groups):
        add_heading(doc, heading, 2)
        add_para(doc, text)
        for key, cap in figs:
            add_image(doc, diagrams[key], cap, 5.95)
        if idx in {1, 3}:
            page_break(doc)
    page_break(doc)


def config_process(doc: Document, diagrams: dict[str, Path]):
    add_heading(doc, "四、SSM 框架配置过程", 1)
    add_para(
        doc,
        "本项目采用 Spring Boot 方式整合 SSM，减少外部 Tomcat 部署步骤。首先在 Maven 中引入 Spring Web、Undertow、Thymeleaf、Validation、MyBatis Starter 和 MySQL 驱动；然后在 application.yml 中配置端口、数据源、连接池、MyBatis XML 位置、实体别名和静态资源压缩；最后通过 WebConfig 注册登录拦截器、CSRF 拦截器和商品图片上传目录映射。",
    )
    add_image(doc, diagrams["config"], "图 14 SSM 框架配置与启动流程图", 5.95)
    add_table(
        doc,
        ["配置项", "配置内容", "作用"],
        [
            ["Maven", "spring-boot-starter-web、undertow、thymeleaf、mybatis、mysql", "提供 Web、模板、持久化和数据库连接能力。"],
            ["数据源", "jdbc:mysql://localhost:3306/ssm_shop，HikariCP 最大连接数 8", "保证数据库连接稳定并控制资源占用。"],
            ["MyBatis", "mapper-locations 与 type-aliases-package", "定位 XML SQL，完成实体映射。"],
            ["拦截器", "AuthInterceptor、CsrfInterceptor", "保护登录态页面和表单提交安全。"],
            ["初始化", "DataSeeder 执行建表、演示数据和幂等迁移", "首次运行准备数据，后续启动只做检查。"],
        ],
        [1.2, 3.3, 2.05],
    )
    add_para(
        doc,
        "配置完成后，项目可通过 Maven 运行或打包为 jar 运行。系统启动时检查 app_meta 初始化标记，未初始化时执行 schema.sql 和 data.sql，已初始化时只做必要迁移检查，避免每次启动重复导入数据。",
    )
    page_break(doc)


def implementation(doc: Document, diagrams: dict[str, Path]):
    add_heading(doc, "五、系统实现与页面展示", 1)
    add_heading(doc, "5.1 前台页面实现", 2)
    add_para(
        doc,
        "前台页面使用统一布局片段组织导航、搜索框、登录状态和消息提醒。首页突出推荐商品、分类入口和公告预告；商品列表提供分类、关键词、价格和排序条件；详情页提供加入购物车、立即购买和收藏入口。页面以浅色背景和绿色主按钮为主，整体更接近真实商城的操作界面。",
    )
    add_image(doc, diagrams["front_board"], "图 15 前台商城页面展示", 6.2)
    add_heading(doc, "5.2 用户中心与运营功能实现", 2)
    add_para(
        doc,
        "用户登录后可以进入活动、背包、邮件、订单和客服页面。活动页展示金币、优惠券和限时活动，背包页用于管理已获得的优惠券，订单页展示交易状态并衔接售后沟通。运营功能会直接影响用户可领取的奖励和结算时可使用的优惠。",
    )
    add_image(doc, diagrams["user_board"], "图 16 用户中心页面展示", 6.2)
    add_heading(doc, "5.3 后台管理实现", 2)
    add_para(
        doc,
        "后台界面围绕日常维护展开。仪表盘显示商品、订单、用户和销售概览；商品管理支持新增、编辑、上下架和删除；活动运营页面可以配置公开活动、站内信奖励和兑换码；订单和客服页面用于处理发货、退款和用户留言。后台与前台共用 Service 和 Mapper，保证业务规则一致。",
    )
    add_image(doc, diagrams["admin_board"], "图 17 后台运营页面展示", 6.2)


def summary(doc: Document):
    add_heading(doc, "六、总结与心得", 1)
    add_para(
        doc,
        "本次课程设计完成了一个较完整的网上商城系统。项目从商品浏览、登录注册和购物车开始，逐步扩展到订单、售后、公告、活动、背包、站内信和后台运营管理。通过这些功能的实现，我对 SSM 分层开发有了更具体的认识：页面负责展示和提交，Controller 接收请求，Service 处理业务规则，Mapper 负责数据访问。只有把职责分清楚，后续功能增加时系统才不会变得混乱。",
    )
    add_para(
        doc,
        "开发过程中比较有收获的是交易流程和运营流程的设计。下单不是简单写入订单记录，还要考虑库存、优惠、金币、地址、支付确认和订单状态；活动奖励也需要判断有效期、领取次数、奖励类型和用户背包。后台管理看起来只是表格操作，但它直接影响前台展示和用户体验，因此更需要保证数据校验和状态流转清楚。",
    )
    add_para(
        doc,
        "当然，系统仍有继续完善的空间，例如可以加入更严格的权限模型、真实支付接口、物流接口和单元测试。总体来说，本项目让我把 JavaEE 课程中学到的控制器、模板、数据库映射、表单校验和系统配置串联起来，也让我认识到项目文档不应只罗列功能，而要说明需求、结构、流程和实现之间的关系。",
    )


def validate_docx(path: Path):
    if not zipfile.is_zipfile(path):
        raise RuntimeError("DOCX is not a valid zip package")
    doc = Document(path)
    text = "\n".join([p.text for p in doc.paragraphs] + [c.text for t in doc.tables for r in t.rows for c in r.cells])
    nonspace = len(re.sub(r"\s+", "", text))
    forbidden = ["重制", "AI", "生成日期", "参考资料", "当前项目截图", "上次", "Codex", "OpenAI", "ChatGPT", "Mermaid", "arc42", "我参考"]
    hits = {word: text.count(word) for word in forbidden if text.count(word)}
    print(f"output={path}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)} nonspace={nonspace}")
    print(f"forbidden_hits={hits}")
    if not (3000 <= nonspace <= 4200):
        raise RuntimeError(f"nonspace count out of requested band: {nonspace}")
    if hits:
        raise RuntimeError(f"forbidden wording found: {hits}")


def main():
    DOCS.mkdir(exist_ok=True)
    diagrams = create_diagrams()
    doc = setup_document()
    cover(doc)
    front_matter(doc)
    requirements_and_design(doc, diagrams)
    detailed_design(doc, diagrams)
    config_process(doc, diagrams)
    implementation(doc, diagrams)
    summary(doc)
    doc.save(OUT)
    validate_docx(OUT)


if __name__ == "__main__":
    main()
