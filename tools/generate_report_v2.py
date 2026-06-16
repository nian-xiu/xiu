from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re
import zipfile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "report-v2-assets"
DIAGRAMS = ASSETS / "diagrams"
OUT = DOCS / "心怡商城项目报告_重制版.docx"

FONT_NAME = "Microsoft YaHei"
FONT_PATH = next(
    (p for p in [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ] if p.exists()),
    None,
)

INK = "#1f2933"
MUTED = "#667085"
TEAL = "#17695c"
TEAL_DARK = "#0f4c43"
BLUE = "#2f6f9f"
AMBER = "#b7791f"
LINE = "#d8e0e5"
SOFT = "#f6f8f7"


def pil_font(size: int, bold: bool = False):
    if FONT_PATH:
        return ImageFont.truetype(str(FONT_PATH), size=size, index=0)
    return ImageFont.load_default()


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str, style: str | None = None, *, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(5)
    if not p.runs:
        r = p.add_run(text)
    else:
        p.runs[0].text = text
        r = p.runs[0]
    set_run_font(r, size=15 if level == 1 else 12, bold=True, color="17483F" if level == 1 else "285B63")
    return p


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def setup_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.68)
    sec.left_margin = Inches(0.78)
    sec.right_margin = Inches(0.78)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in [
        ("Title", 22, "173F4F"),
        ("Heading 1", 15, "173F4F"),
        ("Heading 2", 12, "285B63"),
        ("Heading 3", 11, "6D4C15"),
    ]:
        st = styles[name]
        st.font.name = FONT_NAME
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("心怡商城网上商城系统项目报告")
    set_run_font(run, size=9, color="777777")
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("重制版 · 当前项目截图与精简设计说明")
    set_run_font(run, size=9, color="777777")
    return doc


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, *, bold: bool = False, size: float = 9.2):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, text in enumerate(headers):
        shade_cell(table.rows[0].cells[i], "EAF3F1")
        set_cell(table.rows[0].cells[i], text, bold=True)
        table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell(cells[i], text)
            cells[i].width = Inches(widths[i])
    return table


def caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=9, color="666666")


def add_image(doc: Document, image: Path, cap: str, width: float = 6.55):
    doc.add_picture(str(image), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap)


def wrap(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        if draw.textlength(current + char, font=font) <= width:
            current += char
        else:
            if current:
                lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def rounded(draw, xy, fill, outline=LINE, radius=18, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_arrow(draw, start, end, color=TEAL, width=4):
    x1, y1 = start
    x2, y2 = end
    draw.line([x1, y1, x2, y2], fill=color, width=width)
    direction = 1 if x2 >= x1 else -1
    draw.polygon([(x2, y2), (x2 - direction * 14, y2 - 8), (x2 - direction * 14, y2 + 8)], fill=color)


def diagram_box(draw, xy, title, sub, fill="#ffffff", accent=TEAL):
    rounded(draw, xy, fill=fill, outline="#cdd8dc", radius=18, width=2)
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle([x1, y1, x1 + 10, y2], radius=4, fill=accent)
    title_font = pil_font(25, True)
    sub_font = pil_font(19)
    draw.text((x1 + 28, y1 + 18), title, fill=INK, font=title_font)
    for idx, line in enumerate(wrap(draw, sub, sub_font, x2 - x1 - 48)[:2]):
        draw.text((x1 + 28, y1 + 52 + idx * 25), line, fill=MUTED, font=sub_font)


def make_architecture_diagram():
    img = Image.new("RGB", (1600, 900), "#ffffff")
    d = ImageDraw.Draw(img)
    title = pil_font(38, True)
    d.text((70, 52), "心怡商城系统架构视图", fill=TEAL_DARK, font=title)
    d.text((70, 105), "参考 C4 分层表达：用户界面、Web 控制、业务服务、数据访问与运行基础设施", fill=MUTED, font=pil_font(22))
    boxes = [
        (70, 190, 380, 335, "浏览器 / Thymeleaf", "首页、商品、购物车、订单、后台页面", BLUE),
        (485, 190, 795, 335, "Controller", "路由映射、表单校验、视图模型", TEAL),
        (900, 190, 1210, 335, "Service", "业务规则、事务、权限相关逻辑", AMBER),
        (485, 470, 795, 615, "MyBatis Mapper", "XML SQL、动态查询、主从表写入", TEAL),
        (900, 470, 1210, 615, "MySQL", "用户、商品、订单、活动与客服数据", BLUE),
        (1260, 330, 1510, 485, "Undertow + Maven", "内嵌运行、JAR 启动、资源压缩", "#8a5a2b"),
    ]
    for x1, y1, x2, y2, t, s, a in boxes:
        diagram_box(d, [x1, y1, x2, y2], t, s, accent=a)
    draw_arrow(d, (380, 262), (485, 262), BLUE)
    draw_arrow(d, (795, 262), (900, 262), TEAL)
    draw_arrow(d, (1055, 335), (650, 470), AMBER)
    draw_arrow(d, (795, 542), (900, 542), TEAL)
    draw_arrow(d, (1210, 542), (1260, 430), BLUE)
    d.text((84, 735), "设计要点：控制层不承载复杂业务；Service 层统一事务边界；Mapper 层保持 SQL 可追踪；上传图片通过资源映射暴露。", fill=INK, font=pil_font(24))
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    path = DIAGRAMS / "architecture.png"
    img.save(path)
    return path


def make_sequence_diagram(name: str, messages: list[tuple[int, int, str]], filename: str):
    roles = ["页面", "Controller", "Service", "Mapper", "MySQL"]
    img = Image.new("RGB", (1600, 900), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((70, 50), name, fill=TEAL_DARK, font=pil_font(36, True))
    d.text((70, 98), "统一按 页面 → 控制层 → 业务层 → 持久层 → 数据库 的主链路表达", fill=MUTED, font=pil_font(21))
    xs = [160, 475, 790, 1105, 1390]
    top = 175
    bottom = 815
    for x, role in zip(xs, roles):
        rounded(d, [x - 95, top - 45, x + 95, top + 15], fill="#eff6f4", outline="#b8d5ce", radius=16, width=2)
        tw = d.textlength(role, font=pil_font(24, True))
        d.text((x - tw / 2, top - 30), role, fill=TEAL_DARK, font=pil_font(24, True))
        d.line([x, top + 25, x, bottom], fill="#d8e0e5", width=3)

    y = 305
    msg_font = pil_font(20)
    for i, (src, dst, text) in enumerate(messages, 1):
        x1, x2 = xs[src], xs[dst]
        color = TEAL if i % 2 else BLUE
        draw_arrow(d, (x1, y), (x2, y), color=color, width=4)
        label = f"{i}. {text}"
        maxw = abs(x2 - x1) - 24
        lines = wrap(d, label, msg_font, maxw)
        boxw = min(max(max(d.textlength(line, font=msg_font) for line in lines) + 26, 150), maxw + 30)
        boxh = 30 * len(lines) + 12
        bx = (x1 + x2) / 2 - boxw / 2
        by = y - boxh - 12
        rounded(d, [bx, by, bx + boxw, by + boxh], fill="#ffffff", outline="#e1e8eb", radius=8, width=1)
        for j, line in enumerate(lines[:2]):
            d.text((bx + 13, by + 7 + j * 28), line, fill=INK, font=msg_font)
        y += 76
    path = DIAGRAMS / filename
    img.save(path)
    return path


def make_diagrams() -> dict[str, Path]:
    diagrams = {"architecture": make_architecture_diagram()}
    seqs = {
        "catalog": ("商品浏览与筛选时序图", [
            (0, 1, "访问首页、商品列表或搜索入口"),
            (1, 2, "组装 ProductFilter 并读取当前用户"),
            (2, 3, "查询分类、推荐、最新或筛选商品"),
            (3, 4, "执行 products/categories 动态 SQL"),
            (2, 1, "返回商品集合、分类和筛选条件"),
            (1, 0, "渲染当前版本商城页面"),
        ]),
        "cart_order": ("购物车与订单结算时序图", [
            (0, 1, "提交加入购物车、立即购买或结算请求"),
            (1, 2, "校验登录、地址、支付方式与优惠选择"),
            (2, 3, "读取购物车或构建立即购买摘要"),
            (3, 4, "查询商品、地址和优惠券数据"),
            (2, 3, "事务内写订单、明细、扣库存、核销优惠"),
            (2, 1, "返回订单详情跳转结果"),
        ]),
        "reward": ("活动福利与背包时序图", [
            (0, 1, "领取活动、兑换码或站内邮件奖励"),
            (1, 2, "校验活动状态、名额、过期与重复领取"),
            (2, 3, "记录领取并发放金币或优惠券"),
            (3, 4, "写 activity_claims、reward_mails、user_coupons"),
            (2, 1, "返回领取结果与背包数据"),
            (1, 0, "刷新活动、背包或邮件页面"),
        ]),
        "service": ("订单售后与客服消息时序图", [
            (0, 1, "查看订单、申请退款、确认收货或留言"),
            (1, 2, "按当前用户校验订单归属和状态"),
            (2, 3, "读取订单详情或创建客服消息"),
            (3, 4, "访问 orders、order_items、service_messages"),
            (2, 1, "返回状态变更或消息结果"),
            (1, 0, "刷新订单详情与客服会话"),
        ]),
        "admin": ("后台运营管理时序图", [
            (0, 1, "管理员维护商品、订单、用户、公告和活动"),
            (1, 2, "AuthInterceptor 校验 ADMIN 权限"),
            (2, 3, "调用对应业务服务执行业务校验"),
            (3, 4, "更新商品、订单、公告、用户或运营表"),
            (2, 1, "返回后台列表与仪表盘统计"),
            (1, 0, "渲染运营中枢页面"),
        ]),
        "config": ("SSM 配置与启动时序图", [
            (0, 1, "执行 JAR 启动或 run-app.ps1"),
            (1, 2, "Spring Boot 装配 MVC、拦截器和数据源"),
            (2, 3, "扫描 Mapper 接口并加载 XML"),
            (3, 4, "DataSeeder 检查并初始化表结构"),
            (2, 1, "Undertow 监听 8080 并映射上传目录"),
            (1, 0, "浏览器访问商城首页"),
        ]),
    }
    for key, (title, msgs) in seqs.items():
        diagrams[key] = make_sequence_diagram(title, msgs, f"seq-{key}.png")
    return diagrams


def build_report():
    diagrams = make_diagrams()
    doc = setup_doc()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(130)
    r = p.add_run("心怡商城网上商城系统项目报告")
    set_run_font(r, size=23, bold=True, color="173F4F")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于 Spring MVC + Spring + MyBatis 的 SSM 电商系统")
    set_run_font(r, size=13, color="555555")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("重制版：当前项目截图 · 精简篇幅 · 统一图表风格")
    set_run_font(r, size=11, color="777777")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"生成日期：{datetime.now():%Y-%m-%d}")
    set_run_font(r, size=10, color="888888")

    add_page_break(doc)
    add_heading(doc, "摘 要", 1)
    add_para(doc, "心怡商城是一套面向课程设计与项目展示的网上商城系统，当前版本已从早期商品展示页面迭代为包含前台购物、会员运营、客服沟通和后台管理的完整电商项目。系统采用 Spring Boot 组织工程，在架构上延续 SSM 分层思想：Spring MVC 负责请求映射与视图渲染，Spring Service 负责业务规则和事务边界，MyBatis XML Mapper 负责数据访问，MySQL 负责持久化存储，Undertow 作为内嵌 Web 容器运行。")
    add_para(doc, "本报告按照毕业设计论文常用框架进行压缩整理，重点说明需求分析、系统设计、数据库与配置、核心功能详细设计、当前版本界面实现、测试运行和总结心得。为避免篇幅失控，报告将相近功能合并为核心业务流程，并用统一风格时序图展示调用链。")
    add_para(doc, "关键词：SSM；网上商城；Spring Boot；MyBatis；Thymeleaf；订单事务")
    add_heading(doc, "目 录", 1)
    for line in ["一、需求分析", "二、系统设计", "三、数据库与 SSM 配置", "四、详细设计与核心时序图", "五、系统实现与当前版本截图", "六、测试运行、总结与心得", "参考资料"]:
        add_para(doc, line)

    add_page_break(doc)
    add_heading(doc, "一、需求分析", 1)
    add_para(doc, "心怡商城的目标不是单纯展示商品，而是模拟真实 B2C 商城从浏览、选择、结算、履约到运营维护的闭环流程。普通用户需要能够快速找到商品、加入购物车、选择收货地址、使用金币或优惠券抵扣、提交订单，并在订单产生后查看状态、申请退款、确认收货或联系在线客服。管理员需要通过后台掌握商品、订单、会员、公告与活动的运行状态，及时处理库存、履约和运营问题。")
    add_para(doc, "当前版本相较早期页面增加了运营体验：前台首页包含品牌入口、分类轮播、公告预告、签到金币、精选推荐与新近入库；用户中心包含订单、客服、活动、背包和站内信；后台包含运营仪表盘、商品管理、订单管理、客服消息、活动管理、公告管理和用户管理。系统的主要价值在于把课程项目中常见的登录、增删改查、动态 SQL、事务控制和权限拦截串成一条可演示的业务链。")
    add_table(doc, ["角色", "核心诉求", "主要功能"], [
        ["游客", "快速了解商品与商城内容", "首页、分类、搜索、商品详情、公告查看"],
        ["普通用户", "完成购买、优惠和售后流程", "登录注册、购物车、地址、订单、客服、活动、背包、邮件"],
        ["管理员", "维护商城运营数据", "仪表盘、商品、订单、用户、公告、活动和客服管理"],
    ], [1.0, 2.1, 3.4])

    add_page_break(doc)
    add_heading(doc, "二、系统设计", 1)
    add_para(doc, "参考 arc42 对架构沟通、构件视图和运行视图的划分，以及 C4 模型由高到低描述系统的思路，本项目报告不再堆叠过多无关模块，而是围绕关键层次展开。系统前端由 Thymeleaf 模板与静态 CSS/JS 组成，后端通过 Controller、Service、Mapper 分层组织，数据库通过 schema.sql 与 data.sql 提供结构和演示数据。")
    add_para(doc, "Controller 层只处理路由、参数、表单校验和页面模型；Service 层承担商品筛选、购物车合并、下单事务、优惠核销、订单状态流转、活动领取和后台运营规则；Mapper 层将 SQL 集中到 XML 文件中，便于维护复杂查询和主从表写入。WebConfig 注册登录拦截器和 CSRF 拦截器，保证用户中心、结算、订单、后台路径受到基础保护。")
    add_para(doc, "这种分层方式的好处是职责边界清楚：页面变化一般只影响模板和少量 Controller；业务规则变化主要集中在 Service；SQL 调整则留在 Mapper XML 中。对于课程项目来说，它既能体现传统 SSM 的结构，又能借助 Spring Boot 减少繁琐配置，使项目更容易启动、演示和复现。")
    add_image(doc, diagrams["architecture"], "图 2-1 系统架构视图", width=6.55)

    add_page_break(doc)
    add_heading(doc, "三、数据库与 SSM 配置", 1)
    add_para(doc, "数据库围绕电商核心对象设计，主要表包括 users、categories、products、cart_items、favorites、product_blacklist、addresses、orders、order_items、service_messages、announcements、activity_campaigns、activity_claims、reward_mails、coupon_codes 和 user_coupons。orders 与 order_items 构成主从关系；用户与地址、购物车、收藏、订单、优惠券、站内信形成一对多关系。")
    add_para(doc, "其中 products 保存价格、库存、销量、评分和封面路径；cart_items 以用户和商品建立唯一约束，避免同一用户重复创建同一商品购物车行；orders 保存订单号、用户、地址快照、原价、实付金额、金币抵扣、优惠券抵扣、支付方式和状态；order_items 保存下单时的商品名称、封面、单价、数量和小计，避免后续商品改价影响历史订单。")
    add_para(doc, "SSM 配置集中在 pom.xml、application.yml、启动类和 WebConfig 中。pom.xml 引入 Spring MVC、Thymeleaf、Validation、MyBatis、MySQL Connector 和 Undertow，并排除默认 Tomcat；application.yml 配置 8080 端口、MySQL 数据源、HikariCP 连接池、MyBatis XML 路径、驼峰映射、上传目录和日志等级；启动类通过 @MapperScan 扫描 Mapper；DataSeeder 在首次启动时执行建表和演示数据初始化。")
    add_para(doc, "权限与安全配置采用轻量实现。AuthInterceptor 从 Session 中读取 currentUser，未登录用户访问购物车、结算、订单、地址、客服、活动、背包、设置或后台时会被重定向到登录页；访问 /admin 路径时额外判断用户角色是否为 ADMIN。CsrfInterceptor 为非安全方法校验 _csrf 参数，降低表单被跨站提交的风险。")
    add_table(doc, ["配置项", "位置", "作用"], [
        ["Maven 依赖", "pom.xml", "管理 Spring、MyBatis、MySQL、Undertow 等依赖"],
        ["数据源", "application.yml", "连接 ssm_shop 数据库并配置连接池"],
        ["Mapper 扫描", "SsmShopApplication", "扫描 com.example.ssmshop.mapper 接口"],
        ["拦截器", "WebConfig", "配置登录校验、后台权限、CSRF 和上传资源映射"],
        ["初始化", "DataSeeder + db/*.sql", "首次启动建表、写入商品与演示账号"],
    ], [1.3, 2.0, 3.2])

    add_page_break(doc)
    add_heading(doc, "四、详细设计与核心时序图", 1)
    add_para(doc, "详细设计按“功能组”呈现，避免把每个按钮都拆成重复图。商品浏览、购物车下单、福利背包、售后客服、后台运营和配置启动是本系统最能体现业务完整度的六条主线。时序图参考 Mermaid 对参与者、消息和顺序关系的表达方式绘制，统一使用页面、Controller、Service、Mapper、MySQL 五类参与者。")
    add_heading(doc, "4.1 商品浏览与筛选", 2)
    add_para(doc, "商品列表通过 ProductFilter 承载关键词、分类、价格区间、排序和用户 ID。已登录用户的搜索结果还会结合黑名单过滤，避免被屏蔽商品继续出现在列表。")
    add_image(doc, diagrams["catalog"], "图 4-1 商品浏览与筛选时序图", width=6.55)

    add_page_break(doc)
    add_heading(doc, "4.2 购物车与订单结算", 2)
    add_para(doc, "购物车添加时先校验商品状态和库存，已有商品则合并数量；结算时 OrderService 在事务中完成订单创建、明细写入、扣库存、增加销量、金币扣减、优惠券核销和购物车清空。这个流程是系统数据一致性的核心。")
    add_para(doc, "结算流程中，地址优先使用用户指定地址，未指定时读取默认地址；优惠方式在金币和优惠券之间二选一，金币按整百抵扣，优惠券需要检查未使用、未过期和满减门槛。任何一个环节失败都会抛出异常并触发事务回滚，避免出现订单生成成功但库存未扣、优惠券已核销但订单失败等问题。")
    add_image(doc, diagrams["cart_order"], "图 4-2 购物车与订单结算时序图", width=6.55)
    add_heading(doc, "4.3 活动福利与背包", 2)
    add_para(doc, "活动、邮件和兑换码都统一进入 RewardCenterService 校验，包括状态、领取窗口、名额、过期时间和重复领取。奖励最终转化为金币或用户优惠券，供结算页抵扣使用。")
    add_image(doc, diagrams["reward"], "图 4-3 活动福利与背包时序图", width=6.55)

    add_page_break(doc)
    add_heading(doc, "4.4 订单售后与客服消息", 2)
    add_para(doc, "订单详情按当前用户读取，退款申请与确认收货均需要检查订单状态。客服消息支持用户留言和管理员回复，并通过未读标记在导航栏和后台页面提示。")
    add_image(doc, diagrams["service"], "图 4-4 订单售后与客服消息时序图", width=6.55)
    add_heading(doc, "4.5 后台运营管理", 2)
    add_para(doc, "后台入口统一走 /admin 路径，AuthInterceptor 负责判断 ADMIN 角色。管理员可维护商品、订单、用户、公告、活动、邮件、兑换码和客服消息，仪表盘聚合关键运营指标。")
    add_para(doc, "后台订单状态流转不是任意修改，而是根据当前状态限制目标状态，例如已支付订单可以发货、申请退款或取消，已发货订单可以完成或申请退款，已完成、已退款和已取消订单不能继续回退。这样能让后台演示更接近真实业务。")
    add_image(doc, diagrams["admin"], "图 4-5 后台运营管理时序图", width=6.55)

    add_page_break(doc)
    add_heading(doc, "4.6 配置启动流程", 2)
    add_para(doc, "项目使用 JAR 方式运行，启动时完成 Spring Boot 自动装配、Mapper 扫描、数据初始化、上传目录映射和 Undertow 端口监听。Windows 环境中推荐指定 java.io.tmpdir 到项目 tmp 目录，避免 Undertow 临时目录权限问题。")
    add_image(doc, diagrams["config"], "图 4-6 SSM 配置与启动时序图", width=6.55)

    add_page_break(doc)
    add_heading(doc, "五、系统实现与当前版本截图", 1)
    add_para(doc, "本章截图全部由当前项目在 localhost:8080 实时运行后采集，不再沿用旧版本素材。截图采用首屏或关键区域，避免整页长图造成 Word 中细节不可读。")
    add_image(doc, ASSETS / "current-home.png", "图 5-1 当前版本商城首页", width=6.55)

    add_page_break(doc)
    add_para(doc, "商品列表页展示当前商品图片、分类、库存状态、销量、评分和价格信息，能够直观看到商品搜索、分类筛选和排序入口。")
    add_image(doc, ASSETS / "current-products.png", "图 5-2 当前版本商品列表页", width=6.55)

    add_page_break(doc)
    add_para(doc, "活动中心体现当前版本新增的会员运营能力，展示金币规则、活动数量、限时活动、优惠券奖励和领取状态。")
    add_image(doc, ASSETS / "current-activity.png", "图 5-3 当前版本活动中心页", width=6.55)

    add_page_break(doc)
    add_para(doc, "后台仪表盘以运营视角组织信息，包括在售商品、低库存、待履约、成交额和订单结构，方便管理员快速进入库存、订单和活动管理。")
    add_image(doc, ASSETS / "current-admin-dashboard.png", "图 5-4 当前版本后台仪表盘", width=6.55)

    add_page_break(doc)
    add_para(doc, "活动管理页用于创建公开活动、发送奖励邮件和生成兑换码，是当前版本区别于普通商城 CRUD 项目的重要运营模块。")
    add_image(doc, ASSETS / "current-admin-rewards.png", "图 5-5 当前版本后台活动管理页", width=6.55)

    add_page_break(doc)
    add_heading(doc, "六、测试运行、总结与心得", 1)
    add_para(doc, "运行环境需要 JDK 21、Maven、MySQL 8.0 或 9.x。启动前创建 ssm_shop 数据库，默认连接账号为 root / 123456，可在 application.yml 中修改。构建命令为 mvn package -DskipTests，运行命令可使用 java -Djava.io.tmpdir=tmp -jar target/ssm-shop-1.0.0.jar，也可执行 tools/run-app.ps1。启动后访问 http://localhost:8080/ 进入商城首页，普通用户账号为 customer / 123456，管理员账号为 admin / admin123。")
    add_para(doc, "本次报告重制也修正了素材采集方式：不再复用旧图，而是先启动当前 JAR，再通过独立 Chrome 调试会话访问 localhost 页面，分别以游客、普通用户和管理员身份抓取截图。这样可以保证报告中的首页、活动中心和后台仪表盘与当前代码版本一致。")
    add_para(doc, "测试重点包括：未登录访问购物车、订单和后台时是否被拦截；商品筛选、收藏、黑名单和购物车数量是否正确；下单后库存、销量、订单明细、优惠券和购物车是否同步变化；订单退款、收货和后台状态流转是否合法；活动领取、邮件奖励和兑换码是否能防止重复领取；后台用户启停后登录结果是否符合预期。")
    add_para(doc, "通过本项目可以体会到，优秀的课程项目报告不应只堆功能列表，而应讲清楚系统为什么这样设计、关键流程如何保证数据一致、截图是否来自真实运行版本、图表是否能帮助读者快速理解架构。对我而言，本项目最大的收获是把 SSM 分层、动态 SQL、事务控制、拦截器、模板渲染和本地部署串成了一个可运行、可展示、可写入简历的完整项目。后续若继续优化，可加入接口化测试、分页、评价、支付沙箱、日志审计和更严格的密码加密机制。")

    add_heading(doc, "参考资料", 1)
    add_para(doc, "1. arc42 Template：用于参考架构文档的目标、约束、构件视图、运行视图和部署视图组织方式。")
    add_para(doc, "2. C4 Model：用于参考软件系统、容器、组件等分层架构图表达方式。")
    add_para(doc, "3. Mermaid Sequence Diagram：用于参考时序图参与者、消息顺序和流程表达规范。")
    add_para(doc, "4. Microsoft Word Captions：用于参考 Word 中图片与表格题注的处理方式。")

    DOCS.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def validate(path: Path):
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        for item in ["[Content_Types].xml", "word/document.xml", "word/styles.xml"]:
            if item not in names:
                raise RuntimeError(f"缺少 DOCX 必需文件: {item}")
    text = extract_docx_text(path)
    cjk_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    if not 3000 <= cjk_count <= 4000:
        raise RuntimeError(f"中文字符数不在 3000-4000 范围内: {cjk_count}")
    print(path)
    print(f"中文字符数: {cjk_count}")


if __name__ == "__main__":
    report = build_report()
    validate(report)
