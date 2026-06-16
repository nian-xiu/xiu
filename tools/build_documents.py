from pathlib import Path
from datetime import datetime
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "assets"
OUT = ROOT / "docs"

FONT_PATHS = [
    Path("C:/Windows/Fonts/msyh.ttc"),
    Path("C:/Windows/Fonts/simhei.ttf"),
    Path("C:/Windows/Fonts/simsun.ttc"),
]
FONT_PATH = next((p for p in FONT_PATHS if p.exists()), None)


def font(size, bold=False):
    if FONT_PATH:
        return ImageFont.truetype(str(FONT_PATH), size, index=0)
    return ImageFont.load_default()


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=9.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def apply_table_style(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            for margin_name in ["top", "bottom", "start", "end"]:
                mar = tc_pr.find(qn(f"w:tcMar"))
                if mar is None:
                    mar = OxmlElement("w:tcMar")
                    tc_pr.append(mar)
                node = mar.find(qn(f"w:{margin_name}"))
                if node is None:
                    node = OxmlElement(f"w:{margin_name}")
                    mar.append(node)
                node.set(qn("w:w"), "120")
                node.set(qn("w:type"), "dxa")


def setup_doc(title):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 22, "17483F"),
        ("Heading 1", 16, "17483F"),
        ("Heading 2", 13, "2E5E57"),
        ("Heading 3", 11.5, "7A4F1D"),
    ]:
        st = styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True

    header = sec.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color="777777")
    footer = sec.footer.paragraphs[0]
    footer.text = "心怡商城课程设计文档"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=9, color="777777")
    return doc


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, color="666666")


def add_image(doc, path, caption, width=6.3):
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def draw_sequence(name, participants, messages, filename):
    width, height = 1500, 900
    img = Image.new("RGB", (width, height), "#fbfaf7")
    d = ImageDraw.Draw(img)
    title_font = font(34)
    text_font = font(24)
    small_font = font(20)
    line_color = "#33534d"
    accent = "#1f5f54"
    gold = "#b08a4b"

    d.rounded_rectangle([30, 24, width - 30, height - 24], radius=26, outline="#d8c8ad", width=4, fill="#fffdfa")
    d.text((70, 54), name, fill="#17483f", font=title_font)
    top = 135
    bottom = height - 70
    xs = [110 + i * ((width - 220) / (len(participants) - 1)) for i in range(len(participants))]

    for x, label in zip(xs, participants):
        d.rounded_rectangle([x - 90, top - 46, x + 90, top + 10], radius=16, fill=accent, outline="#17483f", width=2)
        tw = d.textlength(label, font=small_font)
        d.text((x - tw / 2, top - 32), label, fill="white", font=small_font)
        d.line([x, top + 14, x, bottom], fill="#c6b99e", width=3)

    y = top + 72
    step_gap = 60
    for idx, (src, dst, label) in enumerate(messages, start=1):
        x1, x2 = xs[src], xs[dst]
        color = accent if idx % 2 else gold
        d.line([x1, y, x2, y], fill=color, width=4)
        arrow = 12 if x2 > x1 else -12
        d.polygon([(x2, y), (x2 - arrow, y - 8), (x2 - arrow, y + 8)], fill=color)
        msg = f"{idx}. {label}"
        tx = min(x1, x2) + abs(x2 - x1) / 2 - d.textlength(msg, font=small_font) / 2
        d.rounded_rectangle([tx - 10, y - 34, tx + d.textlength(msg, font=small_font) + 10, y - 8], radius=8, fill="#fffdfa")
        d.text((tx, y - 34), msg, fill="#25221d", font=small_font)
        y += step_gap

    img.save(ASSETS / filename)


def build_diagrams():
    roles = ["用户/管理员", "浏览器页面", "Controller", "Service", "Mapper", "MySQL"]
    diagrams = [
        ("用户注册与登录时序图", [(0, 1, "提交账号密码"), (1, 2, "POST /login 或 /register"), (2, 3, "校验表单与密码摘要"), (3, 4, "查询/写入用户"), (4, 5, "执行 SQL"), (3, 2, "返回用户对象"), (2, 1, "写入 Session 并跳转")], "seq-auth.png"),
        ("商品浏览与筛选时序图", [(0, 1, "输入关键词/选择分类"), (1, 2, "GET /products"), (2, 3, "组装 ProductFilter"), (3, 4, "调用 search"), (4, 5, "按分类、价格、排序查询"), (5, 4, "返回商品集合"), (2, 1, "渲染商品网格")], "seq-catalog.png"),
        ("商品详情与收藏时序图", [(0, 1, "进入详情页/点击收藏"), (1, 2, "GET /products/{id} 或 POST /favorites/toggle"), (2, 3, "查询商品和收藏状态"), (3, 4, "调用商品/收藏 Mapper"), (4, 5, "查询或增删收藏记录"), (3, 2, "返回状态"), (2, 1, "刷新详情页")], "seq-favorite.png"),
        ("购物车管理时序图", [(0, 1, "加入购物车/更新数量"), (1, 2, "POST /cart/add 或 /update"), (2, 3, "校验商品状态与库存"), (3, 4, "查询已有购物项"), (4, 5, "插入或更新 cart_items"), (3, 2, "返回购物车摘要"), (2, 1, "跳转购物车页")], "seq-cart.png"),
        ("订单结算时序图", [(0, 1, "选择地址并提交订单"), (1, 2, "POST /checkout"), (2, 3, "读取购物车和地址"), (3, 4, "扣减库存/生成订单"), (4, 5, "写 orders 和 order_items"), (5, 4, "提交事务"), (2, 1, "跳转订单详情")], "seq-order.png"),
        ("地址管理时序图", [(0, 1, "新增/编辑/删除地址"), (1, 2, "POST /addresses"), (2, 3, "校验 AddressForm"), (3, 4, "清除默认地址并保存"), (4, 5, "更新 addresses 表"), (3, 2, "返回结果"), (2, 1, "刷新地址列表")], "seq-address.png"),
        ("后台商品管理时序图", [(0, 1, "新增或编辑商品"), (1, 2, "POST /admin/products"), (2, 3, "校验 ProductForm"), (3, 4, "保存商品信息"), (4, 5, "insert/update products"), (5, 4, "返回影响行数"), (2, 1, "跳转商品列表")], "seq-admin-product.png"),
        ("后台订单处理时序图", [(0, 1, "筛选订单/更新状态"), (1, 2, "GET /admin/orders 或 POST /status"), (2, 3, "读取订单或变更状态"), (3, 4, "调用 OrderMapper"), (4, 5, "查询/更新 orders"), (3, 2, "返回订单详情"), (2, 1, "渲染后台页面")], "seq-admin-order.png"),
        ("后台用户管理时序图", [(0, 1, "启用或停用用户"), (1, 2, "POST /admin/users/{id}/status"), (2, 3, "检查管理员权限"), (3, 4, "更新用户状态"), (4, 5, "update users set status"), (3, 2, "返回结果"), (2, 1, "刷新用户列表")], "seq-admin-user.png"),
    ]
    for title, msgs, fn in diagrams:
        draw_sequence(title, roles, msgs, fn)


def add_overview_table(doc):
    table = doc.add_table(rows=1, cols=4)
    apply_table_style(table, [1.3, 1.7, 1.7, 1.8])
    headers = ["角色", "核心目标", "主要功能", "约束"]
    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], "E8EEF5")
        set_cell_text(table.rows[0].cells[i], h, True)
    rows = [
        ["游客", "快速了解商品", "首页、分类、搜索、详情", "不可下单"],
        ["普通用户", "完成购买闭环", "注册登录、收藏、购物车、地址、下单、订单", "需登录"],
        ["管理员", "运营维护商城", "仪表盘、商品、订单、用户管理", "需 ADMIN 权限"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
    return table


def build_report():
    doc = setup_doc("心怡商城课程设计报告")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("心怡商城网上商城系统课程设计报告")
    set_run_font(r, size=22, bold=True, color="17483F")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于 Spring MVC + Spring + MyBatis + Undertow 的 SSM 商城系统")
    set_run_font(r, size=12, color="666666")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"生成日期：{datetime.now():%Y-%m-%d}")
    set_run_font(r, size=10, color="777777")

    doc.add_paragraph("摘  要", style="Heading 1")
    add_para(doc, "本课程设计围绕“心怡商城”网上商城系统展开，系统采用 Spring Boot 作为工程运行与自动配置基础，在业务结构上保持 SSM 的分层思想：Spring MVC 负责请求分发与页面渲染，Spring Service 承担业务编排与事务控制，MyBatis XML Mapper 完成数据库访问。项目不使用外部 Tomcat，而是通过 Undertow 内嵌服务器运行。系统实现了商品展示、搜索筛选、用户登录注册、收藏、购物车、地址管理、下单、订单查看，以及后台商品、订单、用户管理等功能，形成了较完整的商城业务闭环。")
    add_para(doc, "关键词：SSM；Spring MVC；MyBatis；网上商城；课程设计；Undertow")

    doc.add_paragraph("目  录", style="Heading 1")
    for item in ["一、需求分析", "二、系统设计", "三、详细设计与时序图", "四、SSM 框架配置过程", "五、系统实现与页面展示", "六、总结与心得"]:
        add_para(doc, item)

    doc.add_page_break()
    doc.add_paragraph("一、需求分析", style="Heading 1")
    add_para(doc, "心怡商城定位为面向普通消费者的综合型网上商城，商品覆盖数码电器、居家生活、运动户外、食品生鲜、图书文创五类。系统需要支持用户从浏览、筛选、收藏到加入购物车、提交订单的完整流程，同时为管理员提供基础运营后台。与单一商品展示网站相比，本项目更强调多角色、多模块和数据库驱动的业务闭环。")
    add_overview_table(doc)
    add_para(doc, "功能需求包括：用户可注册登录、维护收货地址、收藏商品、管理购物车、提交订单和查看订单；游客可浏览首页、分类专区、商品详情和搜索结果；管理员可查看运营指标，维护商品上下架与库存价格，处理订单状态，启用或停用用户。非功能需求包括：页面结构清晰、数据初始化可重复执行、业务分层明确、运行方式简洁、数据库操作可维护，并保证基本的权限拦截。")

    doc.add_paragraph("二、系统设计", style="Heading 1")
    add_para(doc, "系统采用典型 MVC 分层结构。Controller 层接收浏览器请求并返回 Thymeleaf 页面；Service 层封装登录、购物车、订单、收藏、地址和后台统计等业务逻辑；Mapper 层通过 XML SQL 与 MySQL 交互；domain、dto、form 分别承载数据库实体、页面聚合对象和表单校验对象。这样的结构既符合传统 SSM 项目的设计习惯，也利用 Spring Boot 简化依赖、端口、数据源和 SQL 初始化配置。")
    table = doc.add_table(rows=1, cols=3)
    apply_table_style(table, [1.8, 2.4, 2.3])
    for i, h in enumerate(["层次", "主要类/文件", "职责"]):
        shade_cell(table.rows[0].cells[i], "E8EEF5")
        set_cell_text(table.rows[0].cells[i], h, True)
    for row in [
        ["表示层", "Thymeleaf + static/css/js", "首页、商品页、购物车、后台页面展示"],
        ["控制层", "ShopController、CartController、AdminController", "路由分发、参数接收、模型封装"],
        ["业务层", "CatalogService、OrderService、CartService", "业务规则、事务、状态流转"],
        ["持久层", "Mapper 接口 + mapper/*.xml", "SQL 查询、插入、更新、删除"],
        ["数据层", "MySQL + schema.sql/data.sql", "用户、商品、订单、地址等数据存储"],
    ]:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
    add_para(doc, "数据库设计围绕商城主流程展开：users 存储账号与角色，categories 和 products 存储商品分类及商品信息，cart_items 与 favorites 支持用户行为，addresses 存储收货信息，orders 和 order_items 保存订单主从数据。下单过程通过事务扣减库存、累计销量并清空购物车，保证核心流程的一致性。")

    doc.add_paragraph("三、详细设计与时序图", style="Heading 1")
    add_para(doc, "详细设计重点关注每个功能从页面请求到数据库访问的调用链。系统中多数流程遵循“浏览器页面—Controller—Service—Mapper—MySQL”的顺序，Controller 只负责接收参数与选择视图，Service 负责业务判断，例如库存校验、默认地址处理、订单事务、密码摘要比对和权限状态判断，Mapper 则保持 SQL 明确可追踪。")
    diagram_files = [
        ("用户注册与登录", "seq-auth.png"),
        ("商品浏览与筛选", "seq-catalog.png"),
        ("商品详情与收藏", "seq-favorite.png"),
        ("购物车管理", "seq-cart.png"),
        ("订单结算", "seq-order.png"),
        ("地址管理", "seq-address.png"),
        ("后台商品管理", "seq-admin-product.png"),
        ("后台订单处理", "seq-admin-order.png"),
        ("后台用户管理", "seq-admin-user.png"),
    ]
    for idx, (caption, fn) in enumerate(diagram_files, 1):
        add_image(doc, ASSETS / fn, f"图 3-{idx} {caption}时序图", width=6.35)

    doc.add_paragraph("四、SSM 框架配置过程", style="Heading 1")
    add_para(doc, "项目采用 Maven 管理依赖，核心依赖包括 spring-boot-starter-web、spring-boot-starter-thymeleaf、mybatis-spring-boot-starter、mysql-connector-j 和 spring-boot-starter-undertow。为满足“不使用 Tomcat”的要求，pom.xml 中排除了 spring-boot-starter-tomcat，并引入 Undertow 作为内嵌 Servlet 容器。")
    add_para(doc, "application.yml 统一配置服务端口、MySQL 数据源、MyBatis mapper-locations、type-aliases-package、驼峰映射和 SQL 初始化脚本。启动时 Spring Boot 会执行 schema.sql 和 data.sql，完成建表与演示数据写入。WebConfig 注册 AuthInterceptor，对购物车、订单、地址、收藏和后台路径进行登录检查，对 /admin 路径进一步校验 ADMIN 角色。")
    add_para(doc, "MyBatis 采用接口与 XML 分离方式，Mapper 接口定义方法，XML 文件承载动态 SQL。ProductMapper.search 通过 ProductFilter 实现关键词、分类、价格范围和排序组合查询；OrderMapper 与 OrderItemMapper 支持订单主从表写入；CartMapper、FavoriteMapper、AddressMapper 分别处理用户行为数据。")

    doc.add_paragraph("五、系统实现与页面展示", style="Heading 1")
    add_para(doc, "前端页面使用 Thymeleaf 模板渲染，公共顶部栏抽取为 fragments/layout.html。首页采用简洁偏欧式的视觉风格，并为五大分类制作 3D 轮播入口；商品列表页支持筛选和排序；详情页展示价格、库存、销量、评分并提供收藏和加入购物车操作；后台页面以表格和表单为主，便于管理员维护。")
    for img, cap in [
        ("screenshot-home.png", "图 5-1 商城首页与分类轮播"),
        ("screenshot-products.png", "图 5-2 商品列表与筛选"),
        ("screenshot-detail.png", "图 5-3 商品详情页"),
        ("screenshot-login.png", "图 5-4 登录页面"),
        ("screenshot-admin.png", "图 5-5 后台仪表盘"),
    ]:
        add_image(doc, ASSETS / img, cap, width=6.35)
    add_para(doc, "订单实现是本系统的关键部分。OrderService.createFromCart 首先读取购物车汇总和收货地址，随后创建订单主表记录，遍历购物车行扣减库存、累计销量并写入订单明细，最后清空购物车。整个过程使用 @Transactional 保证失败回滚。后台订单处理则通过状态字段驱动，支持已支付、已发货、已完成和已取消等状态。")

    doc.add_paragraph("六、总结与心得", style="Heading 1")
    add_para(doc, "通过本次课程设计，我对 SSM 分层思想、MVC 请求流程、MyBatis XML 动态 SQL、事务控制和权限拦截有了更直观的理解。相比只完成单表增删改查，商城系统要求把用户、商品、购物车、订单、地址等对象串联起来，任何一个环节设计不清晰都会影响后续流程。")
    add_para(doc, "项目的主要收获在于：第一，分层设计能明显降低复杂度，Controller、Service、Mapper 的职责边界越清楚，代码越容易维护；第二，数据库脚本和演示数据对课程展示非常重要，能够保证项目在新环境中快速运行；第三，页面体验虽不是后端项目的唯一重点，但良好的导航、轮播、筛选和后台表格能显著提升系统完整度。后续若继续完善，可加入支付模拟、订单评价、分页、上传图片、日志审计和更严格的密码加密机制，使系统更接近真实生产场景。")

    path = OUT / "心怡商城课程设计报告.docx"
    doc.save(path)
    return path


def build_resume_doc():
    doc = setup_doc("心怡商城简历项目经历")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("心怡商城项目经历写法")
    set_run_font(r, size=20, bold=True, color="17483F")
    add_para(doc, "以下内容可按个人简历篇幅直接摘用。若你的简历还有其他项目，本项目建议放在“项目经历”中作为一个完整的 Web 全栈/后端课程项目展示。")

    doc.add_paragraph("项目经历（推荐写法）", style="Heading 1")
    add_para(doc, "心怡商城｜基于 SSM 的网上商城系统")
    add_para(doc, "技术栈：Java 21、Spring Boot、Spring MVC、Spring、MyBatis、Thymeleaf、MySQL、Maven、Undertow")
    add_para(doc, "项目描述：独立设计并实现一个综合型网上商城系统，包含商品浏览、分类筛选、商品详情、用户注册登录、收藏、购物车、地址管理、下单、订单查看，以及后台商品、订单、用户管理等功能。项目采用 SSM 分层结构，使用 Spring Boot 简化工程配置，并以 Undertow 替代外部 Tomcat 运行。")
    bullets = [
        "负责项目整体架构设计，按 controller、service、mapper、domain、dto、form、templates 等目录组织代码，保证业务分层清晰。",
        "使用 MyBatis XML 实现商品动态筛选、订单主从表写入、购物车更新、收藏与地址管理等数据库操作。",
        "设计订单结算流程，在 Service 层通过事务完成购物车读取、库存扣减、销量累计、订单明细写入和购物车清空。",
        "实现登录拦截与后台权限控制，对购物车、订单、地址和后台管理路径进行访问保护。",
        "完成 Thymeleaf 页面与商城 UI 优化，包含分类轮播、商品列表、详情页、后台仪表盘和管理表单。",
    ]
    for b in bullets:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        r = p.add_run("• " + b)
        set_run_font(r, size=10.5)

    doc.add_paragraph("简历亮点提炼", style="Heading 1")
    table = doc.add_table(rows=1, cols=2)
    apply_table_style(table, [1.6, 4.8])
    for i, h in enumerate(["亮点", "可表达能力"]):
        shade_cell(table.rows[0].cells[i], "E8EEF5")
        set_cell_text(table.rows[0].cells[i], h, True)
    for row in [
        ["SSM 分层", "体现 Java Web 项目结构设计、MVC 思维和后端分层能力"],
        ["订单事务", "体现对库存、订单、购物车等业务一致性的理解"],
        ["后台管理", "体现管理端 CRUD、状态流转和权限控制经验"],
        ["页面优化", "体现能借助 AI/Vibe Coding 快速完成可展示产品原型"],
    ]:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)

    doc.add_paragraph("面试表达建议", style="Heading 1")
    add_para(doc, "如果面试官追问项目，可以重点讲三点：一是为什么使用 SSM 分层结构；二是订单结算为什么需要事务；三是如何通过拦截器区分普通用户和管理员。若要强调 Vibe Coding，可表述为“在 AI 辅助下完成需求拆分、页面优化和文档生成，但核心业务流程、数据模型与运行验证由本人理解并整合”。")
    add_image(doc, ASSETS / "screenshot-home.png", "项目首页截图", width=5.8)
    path = OUT / "心怡商城简历项目经历.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    build_diagrams()
    print(build_report())
    print(build_resume_doc())
